from __future__ import annotations

import io
import os
import re
import sys
import tempfile
import unittest
import zipfile
from collections import Counter
from pathlib import Path
from unittest import mock

from PIL import Image, ImageChops, ImageDraw, ImageFont
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import mineru_adapter
from handwriting_markdown_renderer import (
    FontCache,
    HandwritingRenderConfig,
    ScriptBox,
    TextBox,
    _layout_inline,
    _text_to_boxes,
    latex_to_debug_text,
    latex_to_box,
    markdown_render_debug_text,
    render_markdown_handwriting,
    write_images_to_docx_bytes,
)
from markdown_math import editable_docx_bytes, inspect_docx_math, normalize_math_markdown
from mineru_adapter import (
    MinerUConfigError,
    MinerUExtractionError,
    extract_pdf_to_markdown,
    sanitize_mineru_markdown,
    user_facing_mineru_error,
)
from source_extract import extract_source_to_markdown, repair_extracted_markdown_text, safe_source_filename


FONT_PATH = Path(__file__).resolve().parents[1] / "font_assets" / "神韵英子楷书.ttf"


def assert_semantic_tokens_preserved(
    case: unittest.TestCase,
    text: str,
    *,
    text_tokens: tuple[str, ...] = (),
    debug_tokens: tuple[str, ...] = (),
    forbidden_pattern: str | None = None,
) -> None:
    compact_text = re.sub(r"\s+", "", text)
    if forbidden_pattern:
        case.assertNotRegex(text, forbidden_pattern)
    for token in text_tokens:
        case.assertIn(token, compact_text)
    for token in debug_tokens:
        case.assertIn(token, text)


class UnifiedHandwritingPipelineTests(unittest.TestCase):
    def test_formula_debug_text_does_not_keep_latex_macros(self):
        expr = (
            r"\textstyle \frac{a_1}{b^2}+\sqrt{x}+\sum_{i=1}^{n}x_i+\ldots+"
            r"\begin{pmatrix}1&2\\3&4\end{pmatrix}"
        )
        text = latex_to_debug_text(expr, FONT_PATH)
        for macro in ("\\frac", "\\sqrt", "\\sum", "\\begin", "textstyle", "ldots", "dots"):
            self.assertNotIn(macro, text)
        self.assertIn("√", text)
        self.assertIn("∑", text)
        self.assertIn("…", text)
        self.assertIn("[[1,2];[3,4]]", text)

    def test_markdown_math_normalizer_repairs_lone_display_delimiter_and_docx_math(self):
        markdown = (
            r"P \left(Y_{n+k}=j_{n+k}, 1 \leq k \leq m / X_n=i\right)"
            "\n$$\n\n"
            r"其中 $x _ { n } = i.$ 右端只含当前状态 $\textstyle X _ { n } = i$ 和未来参数。"
        )
        normalized = normalize_math_markdown(markdown)
        self.assertIn("$$\nP \\left", normalized)
        self.assertIn("$X_{n} = i$", normalized)
        self.assertNotIn("\\textstyle", normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreater(docx_info["office_math_objects"], 0)
        self.assertFalse(docx_info["has_latex_residuals"])

    def test_docx_math_inspector_detects_broad_latex_residuals(self):
        document = Document()
        document.add_paragraph(r"泄漏公式 a\equiv b, \therefore x\ne0")
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "raw_latex.docx"
            document.save(docx)
            docx_info = inspect_docx_math(docx.read_bytes())

        self.assertTrue(docx_info["has_latex_residuals"])

    def test_docx_math_inspector_detects_unknown_latex_residuals(self):
        document = Document()
        document.add_paragraph(r"泄漏公式 \partial f/\partial x 与 \unknowncmd{x}")
        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "unknown_latex.docx"
            document.save(docx)
            docx_info = inspect_docx_math(docx.read_bytes())

        self.assertTrue(docx_info["has_latex_residuals"])

    def test_markdown_normalizer_wraps_common_bare_calculus_line(self):
        normalized = normalize_math_markdown(r"\partial f/\partial x=0")
        self.assertIn("$$\n\\partial f/\\partial x=0\n$$", normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreater(docx_info["office_math_objects"], 0)
        self.assertFalse(docx_info["has_latex_residuals"])

    def test_markdown_normalizer_wraps_bare_math_inside_chinese_text(self):
        normalized = normalize_math_markdown(r"解：\partial f/\partial x=0，所以 a\equiv b\pmod{n}。")
        self.assertIn(r"解：$\partial f/\partial x=0$，所以 $a\equiv b\pmod{n}$。", normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 2)
        self.assertFalse(docx_info["has_latex_residuals"])

    def test_markdown_normalizer_wraps_bare_symbols_functions_and_greek(self):
        normalized = normalize_math_markdown(
            r"已知 a\times b，角度 \sin x，希腊字母 a\alpha b，还有集合 A\cup B。"
        )
        for formula in (r"$a\times b$", r"$\sin x$", r"$a\alpha b$", r"$A\cup B$"):
            self.assertIn(formula, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 4)
        self.assertFalse(docx_info["has_latex_residuals"])

    def test_markdown_normalizer_rewrites_pandoc_unsafe_commands(self):
        markdown = (
            r"角度 a\degree b，斜线 a\diagup b 和 a\diagdown b，"
            r"近似不等式 a\leqsim b 与 a\geqsim b，"
            r"集合差 A\setminus B，小点 a\ldotp b 与 a\cdotp b，"
            r"极限 \injlim_{i=1}^{n} x_i 与 \projlim_{i=1}^{n} x_i，"
            r"隐藏内容 x\hphantom{abc}y，覆盖 \llap{x+y}+\rlap{z}+\smash{w}，"
            r"以及 \limits x+\nolimits y。"
        )
        normalized = normalize_math_markdown(markdown)
        for raw in (
            r"\degree",
            r"\diagup",
            r"\diagdown",
            r"\leqsim",
            r"\geqsim",
            r"\setminus",
            r"\ldotp",
            r"\cdotp",
            r"\injlim",
            r"\projlim",
            r"\hphantom",
            r"\llap",
            r"\rlap",
            r"\smash",
            r"\limits",
            r"\nolimits",
        ):
            self.assertNotIn(raw, normalized)
        for token in ("°", "⟋", "⟍", "⪅", "⪆", "∖", ".", "·", "inj lim", "proj lim", "xy", "x+y", "z", "w"):
            self.assertIn(token, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 8)
        self.assertFalse(docx_info["has_latex_residuals"])

    def test_markdown_normalizer_preserves_inline_display_math_and_image_markers(self):
        markdown = "前文 ![图1](assets/a.png) 中间 $$x^2+1$$ 后文 <img src=\"b.png\" alt=\"图2\">"
        normalized = normalize_math_markdown(markdown)
        self.assertIn("![图1](assets/a.png)", normalized)
        self.assertIn("<img src=\"b.png\" alt=\"图2\">", normalized)
        self.assertIn("$$\nx^2+1\n$$", normalized)
        debug_text = markdown_render_debug_text(normalized, FONT_PATH)
        self.assertIn("[图片:![图1](assets/a.png)]", debug_text)
        self.assertIn("[图片:<img src=\"b.png\" alt=\"图2\">]", debug_text)
        self.assertIn("x^2+1", debug_text)

    def test_markdown_normalizer_decodes_html_entities_before_handwriting_render(self):
        markdown = "不等式 a &lt; b &amp;&amp; c &gt; d，公式 $x &lt; y &amp; y &gt; z$。"
        normalized = normalize_math_markdown(markdown)
        debug_text = markdown_render_debug_text(normalized, FONT_PATH)
        compact_text = re.sub(r"\s+", "", debug_text)
        self.assertIn("a<b&&c>d", compact_text)
        self.assertIn("x<y&y>z", compact_text)
        self.assertNotRegex(debug_text, r"&lt;|&gt;|&amp;")

    def test_unknown_latex_commands_remain_visible_without_raw_latex(self):
        debug_text = latex_to_debug_text(r"\unknowncmd{x}+\overset{a}{b}", FONT_PATH)
        self.assertNotIn("\\", debug_text)
        self.assertIn("unknowncmd", debug_text)
        self.assertIn("x", debug_text)
        self.assertIn("a", debug_text)
        self.assertIn("b", debug_text)

    def test_markdown_normalizer_rewrites_unknown_latex_commands_in_explicit_math(self):
        normalized = normalize_math_markdown(r"未知 $\foo{x}+y$ 结束。")
        self.assertIn(r"$\operatorname{foo}(x)+y$", normalized)
        self.assertNotIn(r"\foo", normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 1)
        self.assertFalse(docx_info["has_latex_residuals"])
        debug_text = markdown_render_debug_text(normalized, FONT_PATH)
        assert_semantic_tokens_preserved(
            self,
            debug_text,
            text_tokens=("未知", "结束"),
            debug_tokens=("foo", "x", "y"),
            forbidden_pattern=r"\\|foo\{x\}",
        )

    def test_markdown_normalizer_preserves_unknown_latex_content_in_plain_text(self):
        markdown = r"未知 \unknowncmd{x}+y 结束，路径 C:\Users\alice 不当公式。"
        normalized = normalize_math_markdown(markdown)
        self.assertNotIn(r"\unknowncmd", normalized)
        self.assertIn("unknowncmd(x)+y", normalized)
        self.assertIn(r"C:\Users\alice", normalized)
        debug_text = markdown_render_debug_text(normalized, FONT_PATH)
        assert_semantic_tokens_preserved(
            self,
            debug_text,
            text_tokens=("未知", "unknowncmd(x)+y", "结束", "C:", "Users", "alice"),
            forbidden_pattern=r"\\unknowncmd|unknowncmd\{x\}",
        )

    def test_literal_backslashes_are_not_dropped_as_unknown_latex(self):
        markdown = r"路径 C:\Users\alice 保留，普通反斜杠 a\b 保留，公式 \frac{1}{2} 转换。"
        normalized = normalize_math_markdown(markdown)
        self.assertIn(r"C:\Users\alice", normalized)
        self.assertIn(r"a\b", normalized)
        self.assertNotIn(r"\frac", markdown_render_debug_text(normalized, FONT_PATH))
        debug_text = markdown_render_debug_text(normalized, FONT_PATH)
        for token in (r"C:\Users\alice", r"a\b", "(1)/(2)"):
            self.assertIn(token, debug_text)

    def test_markdown_renderer_unescapes_plain_text_punctuation_without_dropping_paths(self):
        markdown = r"普通 A\&B，百分号 \%，编号 \#1，下划线 x\_1，路径 C:\Users\alice 保留。"
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("A&B", "%", "#1", "x_1", r"C:\Users\alice"):
            self.assertIn(token, compact_text)
        self.assertNotIn(r"A\&B", debug_text)
        self.assertNotIn(r"\%", debug_text)
        self.assertNotIn(r"\#1", debug_text)

    def test_escaped_plain_text_dollar_does_not_swallow_following_formula(self):
        markdown = r"价格 \$5，斜杠 a\/b，路径 C:\Users\alice，公式 $x_1+y$。"
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("$5", "a/b", r"C:\Users\alice", "x_1+y"):
            self.assertIn(token, compact_text)
        self.assertNotIn(r"\$5", debug_text)
        self.assertNotIn(r"\/", debug_text)
        self.assertNotIn("x_1+y$", debug_text)

    def test_text_like_math_groups_unescape_visible_punctuation(self):
        debug_text = latex_to_debug_text(
            r"\text{A\_B \% done cost \$5 A\&B \#1}+"
            r"\mbox{编号 \#2}+\operatorname{arg\_max}",
            FONT_PATH,
        )
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("A_B", "%done", "cost$5", "A&B", "#1", "编号#2", "arg_max"):
            self.assertIn(token, compact_text)
        self.assertNotRegex(debug_text, r"\\[_%$&#]|text|mbox|operatorname")

    def test_text_like_math_groups_render_latex_space_commands_as_spaces(self):
        debug_text = latex_to_debug_text(
            r"\text{A\ B}+\text{A\quad B}+\mbox{C\qquad D}+\operatorname{arg\,min}",
            FONT_PATH,
        )
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("AB", "CD", "argmin"):
            self.assertIn(token, compact_text)
        self.assertNotRegex(debug_text, r"\\|quad|qquad|operatorname|text|mbox")

    def test_common_math_decorations_have_visible_marks(self):
        debug_text = latex_to_debug_text(r"\overline{x}+\hat{y}+\vec{z}", FONT_PATH)
        self.assertIn("¯x", debug_text)
        self.assertIn("^y", debug_text)
        self.assertIn("→z", debug_text)

    def test_contextual_dots_commands_render_as_ellipsis(self):
        debug_text = latex_to_debug_text(r"\dots+\dotsc+\dotsi+\dotsb+\dotsm+\dotso", FONT_PATH)
        self.assertGreaterEqual(debug_text.count("…"), 6)
        self.assertNotRegex(debug_text, r"\\|dots|ldots")

    def test_infix_over_renders_as_fraction(self):
        self.assertEqual("(a)/(b)", latex_to_debug_text(r"a \over b", FONT_PATH))

    def test_legacy_infix_stack_commands_do_not_render_control_words(self):
        choose_text = latex_to_debug_text(r"n \choose k", FONT_PATH)
        stack_text = latex_to_debug_text(r"a \atop b+c \brack d+e \brace f", FONT_PATH)
        debug_text = choose_text + stack_text
        self.assertNotRegex(debug_text, r"\\|choose|atop|brack|brace")
        self.assertIn("C(n,k)", choose_text)
        for token in ("a", "b", "c", "d", "e", "f"):
            self.assertIn(token, stack_text)

    def test_legacy_dimension_infix_commands_do_not_render_control_words(self):
        above_text = latex_to_debug_text(r"a \above 0pt b", FONT_PATH)
        over_delims_text = latex_to_debug_text(r"c \overwithdelims() d", FONT_PATH)
        atop_delims_text = latex_to_debug_text(r"n \atopwithdelims[] k", FONT_PATH)
        debug_text = above_text + over_delims_text + atop_delims_text
        self.assertNotRegex(debug_text, r"\\|above|overwithdelims|atopwithdelims|0pt")
        for token in ("a", "b", "c", "d", "n", "k"):
            self.assertIn(token, debug_text)

    def test_text_command_preserves_inner_spaces(self):
        self.assertEqual("if x>0", latex_to_debug_text(r"\text{if }x>0", FONT_PATH))

    def test_mathop_and_starred_operatorname_do_not_render_command_names(self):
        mathop_text = latex_to_debug_text(r"\mathop{\lim}\limits_{x\to0} f(x)", FONT_PATH)
        operator_text = latex_to_debug_text(r"\operatorname*{arg\,max}_{x} f(x)", FONT_PATH)
        self.assertNotRegex(mathop_text + operator_text, r"mathop|operatorname|\\")
        self.assertIn("lim_x→0", mathop_text)
        self.assertIn("arg max_x", operator_text)

    def test_optional_root_index_and_extensible_arrow_labels_are_preserved(self):
        debug_text = latex_to_debug_text(
            r"\sqrt[3]{x^2+y}+\xrightarrow[n\to0]{m\to\infty} y+\xleftarrow{k} z",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|sqrt|xrightarrow|xleftarrow")
        for token in ("√[3]", "x^2", "y", "→", "n→0", "m→∞", "←", "k", "z"):
            self.assertIn(token, debug_text)

    def test_common_symbol_variants_render_without_command_words(self):
        debug_text = latex_to_debug_text(
            r"\xleftrightarrow{n\to\infty} y+\xRightarrow{a} b+\xLeftarrow{c} d+"
            r"\varDelta+\varGamma+\overleftrightarrow{CD}+"
            r"a\precsim b+c\succsim d+e\coloneqq f+g\eqqcolon h+i\triangleq j+"
            r"A\leadsto B+A\longmapsto B",
            FONT_PATH,
        )
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("↔", "n→∞", "⇒", "a", "⇐", "c", "Δ", "Γ", "↔CD", "≾", "≿", "≔", "≕", "≜", "↝", "⟼"):
            self.assertIn(token, compact_text)
        self.assertNotRegex(
            debug_text,
            r"\\|xleftrightarrow|xRightarrow|xLeftarrow|varDelta|varGamma|overleftrightarrow|"
            r"precsim|succsim|coloneqq|eqqcolon|triangleq|leadsto|longmapsto",
        )

    def test_common_amssymb_variants_render_without_command_words(self):
        debug_text = latex_to_debug_text(
            r"a\lessdot b+c\gtrdot d+e\lll f+g\ggg h+"
            r"i\nleqslant j+k\ngeqslant l+m\lneqq n+o\gneqq p+"
            r"A\subsetneqq B+C\supsetneqq D+E\varsubsetneq F+G\varsupsetneq H+"
            r"\nexists x+\complement A+\Bbbk+\ulcorner x\urcorner+\llcorner y\lrcorner+"
            r"p\curlywedge q+r\curlyvee s+t\Cap u+v\Cup w+"
            r"x\circledast y+z\circledcirc a+b\circleddash c",
            FONT_PATH,
        )
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in (
            "⋖", "⋗", "⋘", "⋙", "≰", "≱", "≨", "≩", "⫋", "⫌",
            "⊊", "⊋", "∄", "∁", "𝕜", "⌜x⌝", "⌞y⌟", "⋏", "⋎",
            "⋒", "⋓", "⊛", "⊚", "⊝",
        ):
            self.assertIn(token, compact_text)
        self.assertNotRegex(
            debug_text,
            r"\\|lessdot|gtrdot|lll|ggg|nleqslant|ngeqslant|lneqq|gneqq|subsetneqq|supsetneqq|"
            r"varsubsetneq|varsupsetneq|nexists|complement|Bbbk|ulcorner|urcorner|llcorner|lrcorner|"
            r"curlywedge|curlyvee|Cap|Cup|circledast|circledcirc|circleddash",
        )

    def test_triangle_and_misc_symbols_render_without_command_words(self):
        debug_text = latex_to_debug_text(
            r"a\blacktriangleright b+c\blacktriangleleft d+e\trianglerighteq f+"
            r"g\trianglelefteq h+\maltese",
            FONT_PATH,
        )
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("▸", "◂", "⊵", "⊴", "✠"):
            self.assertIn(token, compact_text)
        self.assertNotRegex(
            debug_text,
            r"\\|blacktriangleright|blacktriangleleft|trianglerighteq|trianglelefteq|maltese",
        )

    def test_not_relations_render_as_standard_negated_symbols(self):
        debug_text = latex_to_debug_text(
            r"\not\approx+\not\sim+\not\equiv+\not\leqslant+\not\geqslant+\not\parallel",
            FONT_PATH,
        )
        for token in ("≉", "≁", "≢", "≰", "≱", "∦"):
            self.assertIn(token, debug_text)
        self.assertNotRegex(debug_text, r"¬[≈∼≤≥∥]|\\|not")

    def test_named_delimiter_commands_render_as_visible_delimiters(self):
        debug_text = latex_to_debug_text(
            r"\lbrace x\in A \rbrace+\lparen y \rparen+\lbrack z \rbrack+"
            r"\genfrac{\lbrace}{\rbrace}{0pt}{}{a}{b}",
            FONT_PATH,
        )
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("{x∈A}", "(y)", "[z]", "{(a)/(b)}"):
            self.assertIn(token, compact_text)
        self.assertNotRegex(debug_text, r"\\|lbrace|rbrace|lparen|rparen|lbrack|rbrack|genfrac")

    def test_alignment_environment_and_equation_metadata_do_not_render_control_words(self):
        debug_text = latex_to_debug_text(
            r"\begin{align}a&=b+c\\d&=e+f\tag{1}\label{eq:one}\end{align}",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|begin|end|align|tag|label|&|eq:one")
        for token in ("a", "=b+c", "d", "=e+f", "(1)"):
            self.assertIn(token, debug_text)

    def test_stacked_limit_helpers_do_not_render_control_words_or_alignment_specs(self):
        debug_text = latex_to_debug_text(
            r"\sum_{\substack{i=1\\j=2}}^n a_{ij}+"
            r"\lim_{\begin{subarray}{c}x\to0\\y\to1\end{subarray}} f(x,y)",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|substack|subarray|begin|end|_c")
        for token in ("∑", "i=1", "j=2", "^n", "a", "ij", "lim", "x→0", "y→1", "f"):
            self.assertIn(token, debug_text)

    def test_matrix_variants_render_as_matrix_rows_without_control_words(self):
        debug_text = latex_to_debug_text(
            r"\begin{smallmatrix}1&2\\3&4\end{smallmatrix}+"
            r"\begin{Bmatrix}a&b\\c&d\end{Bmatrix}+"
            r"\begin{Vmatrix}p&q\\r&s\end{Vmatrix}",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|begin|end|smallmatrix|Bmatrix|Vmatrix|&")
        for token in ("[[1,2];[3,4]]", "[[a,b];[c,d]]", "[[p,q];[r,s]]"):
            self.assertIn(token, debug_text)

    def test_matrix_parser_preserves_escaped_ampersands_inside_cells(self):
        debug_text = latex_to_debug_text(
            r"\begin{matrix}a\&b&c\\\text{A&B}&d\end{matrix}+"
            r"\begin{cases}x\&y,&x>0\\0,&x\leq0\end{cases}",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|text|cases|matrix")
        for token in ("a&b", "c", "A&B", "d", "x&y", "x>0", "x≤0"):
            self.assertIn(token, debug_text)

    def test_extended_decorations_and_boxed_content_do_not_render_command_names(self):
        debug_text = latex_to_debug_text(
            r"\overparen{AB}+\underparen{CD}+\overleftarrow{EF}+"
            r"\underleftarrow{GH}+\boxed{a+b}",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|overparen|underparen|overleftarrow|underleftarrow|boxed")
        for token in ("AB", "CD", "EF", "GH", "a+b"):
            self.assertIn(token, debug_text)

    def test_color_cancel_and_middle_commands_preserve_content_without_control_words(self):
        debug_text = latex_to_debug_text(
            r"\color{red}{x+y}+\textcolor{blue}{z}+\cancel{x}+"
            r"\bcancel{y}+\xcancel{z}+\sout{w}+\left\lbrace x\middle|x>0\right\rbrace",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|color|textcolor|cancel|bcancel|xcancel|sout|middle")
        for token in ("x+y", "z", "x", "y", "w", "{", "|", "x>0", "}"):
            self.assertIn(token, debug_text)

    def test_array_rule_and_span_commands_do_not_render_control_words(self):
        debug_text = latex_to_debug_text(
            r"\begin{array}{c|c}\hline a&b\\\cline{1-2}\multicolumn{2}{c}{c+d}\end{array}",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|hline|cline|multicolumn|c\|c")
        for token in ("a", "b", "c+d"):
            self.assertIn(token, debug_text)

    def test_plain_tex_matrix_commands_render_as_structured_matrices(self):
        expr = (
            r"\matrix{a&b\\c&d}+"
            r"\pmatrix{1&2\\3&4}+"
            r"\cases{x^2,&x>0\\0,&x\leq 0}"
        )
        normalized = normalize_math_markdown(f"题目 {expr} 结束")
        self.assertIn(r"\begin{matrix}", normalized)
        self.assertIn(r"\begin{pmatrix}", normalized)
        self.assertIn(r"\begin{cases}", normalized)
        debug_text = markdown_render_debug_text(normalized, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|matrix|pmatrix|cases")
        for token in ("[[a,b];[c,d]]", "[[1,2];[3,4]]", "x^2", "x>0", "x≤0"):
            self.assertIn(token, debug_text)

    def test_direct_plain_tex_matrix_and_stray_structural_commands_do_not_render_names(self):
        debug_text = latex_to_debug_text(
            r"\pmatrix{1&2\\3&4}+\matrix+\subarray+\buildrel x+y+\end{pmatrix}+z",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|pmatrix|matrix|subarray|buildrel|end")
        for token in ("[[1,2];[3,4]]", "x", "y", "z"):
            self.assertIn(token, debug_text)

    def test_named_accent_commands_render_as_decorations(self):
        debug_text = latex_to_debug_text(r"\acute{x}+\grave{y}+\breve{z}+\check{w}+\mathring{A}", FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|acute|grave|breve|check|mathring")
        for token in ("´x", "`y", "˘z", "ˇw", "˚A"):
            self.assertIn(token, debug_text)

    def test_extra_font_wrappers_do_not_render_command_names(self):
        debug_text = latex_to_debug_text(r"\mathscr{F}+\mathds{1}+\bm{x}+\boldsymbol{\alpha}", FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|mathscr|mathds|bm|boldsymbol")
        for token in ("F", "1", "x", "α"):
            self.assertIn(token, debug_text)

    def test_common_named_math_symbols_render_as_symbols(self):
        debug_text = latex_to_debug_text(r"\Re z+\Im z+\ell+\hbar+\aleph+\wp", FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|Re|Im|ell|hbar|aleph|wp")
        for token in ("ℜ", "z", "ℑ", "ℓ", "ℏ", "ℵ", "℘"):
            self.assertIn(token, debug_text)

    def test_logic_relation_and_special_arrow_commands_render_as_symbols(self):
        debug_text = latex_to_debug_text(
            r"\neg p+\implies q+\Longrightarrow r+\Longleftrightarrow s+"
            r"a\prec b+a\preceq b+a\succ b+a\succeq b+a\ll b+a\gg b+"
            r"a\asymp b+a\doteq b+a\hookrightarrow b+a\twoheadrightarrow b+"
            r"a\rightsquigarrow b+A\smallsetminus B+A\sqcup B",
            FONT_PATH,
        )
        self.assertNotRegex(
            debug_text,
            r"\\|neg|implies|Longrightarrow|Longleftrightarrow|prec|preceq|succ|succeq|"
            r"asymp|doteq|hookrightarrow|twoheadrightarrow|rightsquigarrow|smallsetminus|sqcup",
        )
        for token in ("¬", "⇒", "⟹", "⟺", "≺", "≼", "≻", "≽", "≪", "≫", "≍", "≐", "↪", "↠", "↝", "∖", "⊔"):
            self.assertIn(token, debug_text)

    def test_additional_common_latex_symbols_render_readably_without_command_names(self):
        debug_text = latex_to_debug_text(
            r"\varpi+\varsigma+\varrho+a\leftharpoonup b+a\rightleftharpoons b+"
            r"a\leqq b+a\lessapprox b+A\bigcup B+\bullet+\square+\arcsin x+\thinspace y",
            FONT_PATH,
        )
        self.assertNotRegex(
            debug_text,
            r"\\|varpi|varsigma|varrho|leftharpoonup|rightleftharpoons|"
            r"leqq|lessapprox|bigcup|bullet|square|thinspace",
        )
        for token in ("ϖ", "ς", "ϱ", "↼", "⇌", "≦", "⪅", "⋃", "•", "□", "arcsin", "x", "y"):
            self.assertIn(token, debug_text)

    def test_common_ams_symbols_render_readably_without_command_names(self):
        debug_text = latex_to_debug_text(
            r"a\nleq b+a\ngeq b+a\nless b+a\ngtr b+a\lneq b+a\gneq b+"
            r"A\nsubseteq B+A\nsupseteq B+A\Subset B+A\Supset B+"
            r"A\boxtimes B+A\boxplus B+A\ltimes B+A\rightthreetimes B+"
            r"\beth+\Game+\mho+\backprime+a\VDash b+a\nvdash b+a\nVDash b",
            FONT_PATH,
        )
        self.assertNotRegex(
            debug_text,
            r"\\|nleq|ngeq|nless|ngtr|lneq|gneq|nsubseteq|nsupseteq|Subset|Supset|"
            r"boxtimes|boxplus|ltimes|rightthreetimes|beth|Game|mho|backprime|VDash|nvdash|nVDash",
        )
        for token in ("≰", "≱", "≮", "≯", "⪇", "⪈", "⊈", "⊉", "⋐", "⋑", "⊠", "⊞", "⋉", "⋌", "ℶ", "⅁", "℧", "‵", "⊫", "⊬", "⊯"):
            self.assertIn(token, debug_text)

    def test_more_common_symbols_render_readably_without_command_names(self):
        debug_text = latex_to_debug_text(
            r"a\nmid b+a\nparallel b+a\ncong b+a\napprox b+a\nsim b+"
            r"a\nsmile b+a\nfrown b+a\smallsmile b+a\smallfrown b+"
            r"\coprod_{i=1}^{n}A_i+\bigstar+\lozenge+\blacktriangle+"
            r"\clubsuit+\diamondsuit+\heartsuit+\spadesuit+\natural+\flat+\sharp+"
            r"\top+\bot+A\diagup B+A\diagdown B",
            FONT_PATH,
        )
        self.assertNotRegex(
            debug_text,
            r"\\|nmid|nparallel|ncong|napprox|nsim|nsmile|nfrown|smallsmile|smallfrown|"
            r"coprod|bigstar|lozenge|blacktriangle|"
            r"clubsuit|diamondsuit|heartsuit|spadesuit|natural|flat|sharp|top|bot|diagup|diagdown",
        )
        for token in ("∤", "∦", "≇", "≉", "≁", "¬⌣", "¬⌢", "⌣", "⌢", "∐", "★", "◊", "▲", "♣", "♦", "♥", "♠", "♮", "♭", "♯", "⊤", "⊥", "⟋", "⟍"):
            self.assertIn(token, debug_text)

    def test_structural_latex_helpers_preserve_content_without_command_names(self):
        debug_text = latex_to_debug_text(
            r"\stackrel{def}{=}+\genfrac{[}{]}{0pt}{}{a+b}{c+d}+"
            r"\genfrac{\{}{\}}{0pt}{}{n}{k}+"
            r"\mathrel{R}+\smash{x}+\raisebox{1ex}{y}+"
            r"\begin{pmatrix}\hdotsfor{3}\\a&b&c\end{pmatrix}",
            FONT_PATH,
        )
        self.assertNotRegex(
            debug_text,
            r"\\|stackrel|genfrac|mathrel|smash|raisebox|hdotsfor|0pt|1ex",
        )
        for token in ("def", "=", "a+b", "c+d", "n", "k", "R", "x", "y", "⋯", "a", "b", "c"):
            self.assertIn(token, debug_text)

    def test_legacy_buildrel_and_text_style_wrappers_do_not_render_command_names(self):
        debug_text = latex_to_debug_text(
            r"\buildrel def \over =+\buildrel * \over \longrightarrow+"
            r"\pmb{x}+\boldmath{y}+\cal{F}+\Bbb{R}+"
            r"\textnormal{abc}+\textit{def}+\operatornamewithlimits{argmax}_{x}",
            FONT_PATH,
        )
        self.assertNotRegex(
            debug_text,
            r"\\|buildrel|over|longrightarrow|pmb|boldmath|cal|Bbb|"
            r"textnormal|textit|operatornamewithlimits",
        )
        for token in ("def", "=", "*", "⟶", "x", "y", "F", "R", "abc", "argmax"):
            self.assertIn(token, debug_text)

    def test_modular_phantom_and_text_box_helpers_do_not_render_control_words(self):
        debug_text = latex_to_debug_text(
            r"\limsup_{n\to\infty}a_n+\liminf_{n\to\infty}b_n+"
            r"\injlim X+\projlim Y+a\bmod n+a\pod{n}+x\phantom{abc}y+x\mbox{ text }",
            FONT_PATH,
        )
        self.assertNotRegex(debug_text, r"\\|limsup|liminf|injlim|projlim|bmod|pod|phantom|mbox|abc")
        for token in ("lim sup", "lim inf", "inj lim", "proj lim", "mod", "(n)", "xy", "text"):
            self.assertIn(token, debug_text)

    def test_escaped_accent_commands_render_as_decorations(self):
        debug_text = latex_to_debug_text(r"\~{\pi}+\~\pi+\'{e}+\`{a}+\"{u}+x^\pi", FONT_PATH)
        self.assertNotIn("\\", debug_text)
        self.assertNotIn("~(", debug_text)
        self.assertIn("~π", debug_text)
        self.assertIn("´e", debug_text)
        self.assertIn("`a", debug_text)
        self.assertIn("¨u", debug_text)
        self.assertIn("x^π", debug_text)

    def test_common_latex_commands_render_readably_without_raw_latex(self):
        expr = (
            r"\dfrac{a_1}{b^2}+\binom{n}{k}+\overset{a}{b}+"
            r"\underset{0}{\lim}+\operatorname{Var}(X)+\mathbb{R}+"
            r"\iff+\mapsto+\left\{x\mid x\geq0\right\}"
        )
        debug_text = latex_to_debug_text(expr, FONT_PATH)
        self.assertNotIn("\\", debug_text)
        for macro in ("dfrac", "binom", "overset", "underset", "iff", "mapsto", "left", "right"):
            self.assertNotIn(macro, debug_text)
        for token in ("a", "1", "b", "2", "n", "k", "0", "lim", "Var", "X", "R", "x"):
            self.assertIn(token, debug_text)
        self.assertIn("↔", debug_text)
        self.assertIn("↦", debug_text)

    def test_bare_presentation_commands_are_wrapped_as_office_math(self):
        markdown = (
            r"题目 \binom{n}{k}，再看 \overset{a}{b} 和 "
            r"\underset{0}{\lim}，最后 \vec{x}+\dot{y}+\ddot{z}。"
        )
        normalized = normalize_math_markdown(markdown)
        for formula in (
            r"$\binom{n}{k}$",
            r"$\overset{a}{b}$",
            r"$\underset{0}{\lim}$",
            r"$\vec{x}+\dot{y}+\ddot{z}$",
        ):
            self.assertIn(formula, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 4)
        self.assertFalse(docx_info["has_latex_residuals"])

    def test_markdown_debug_text_preserves_body_and_formula_tokens(self):
        markdown = (
            r"第12题：已知 A_n=3，求 $P(X_n=i)=\dfrac{a_1}{b^2}$ 的值。"
            "\n\n$$\n"
            r"f(x)=\begin{cases}x^2+1,&x\geq0\\-x,&x<0\end{cases}"
            "\n$$\n"
            "最后一行 z_9 不可丢。"
        )
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotIn("\\", debug_text)
        for token in ("第12题", "已知", "A", "n", "3", "P", "X", "i", "a", "1", "b", "2", "f", "x", "0", "最后一行", "z", "9"):
            self.assertIn(token, debug_text)

    def test_plain_requests_with_math_use_markdown_renderer(self):
        from app import should_render_with_markdown_renderer

        self.assertTrue(should_render_with_markdown_renderer("plain", r"题目 \dfrac{a}{b}"))
        self.assertTrue(should_render_with_markdown_renderer("plain", r"题目 $x+1$"))
        self.assertTrue(should_render_with_markdown_renderer("plain", r"a\equiv b\pmod{n}"))
        self.assertTrue(should_render_with_markdown_renderer("plain", r"x\perp y,\ \angle ABC"))
        self.assertTrue(should_render_with_markdown_renderer("plain", r"\therefore x\ne 0"))
        self.assertFalse(should_render_with_markdown_renderer("plain", "纯文本内容"))

    def test_legacy_textfileprocess_uses_markdown_formula_extraction(self):
        from app import extract_textfileprocess_content

        with tempfile.TemporaryDirectory() as tmp:
            docx = Path(tmp) / "formula.docx"
            docx.write_bytes(editable_docx_bytes(r"旧入口公式 $\frac{a_1}{b^2}+\sum_{i=1}^{n}x_i$ 完成"))
            result = extract_textfileprocess_content(docx)
        self.assertIn("旧入口公式", result["text"])
        self.assertIn(r"\frac", result["text"])
        self.assertIn(r"\sum", result["text"])
        self.assertIn("完成", result["text"])

    def test_pdf_fallback_text_is_normalized_before_handwriting_render(self):
        import fitz

        from app import extract_textfileprocess_content

        with tempfile.TemporaryDirectory() as tmp:
            pdf = Path(tmp) / "formula.pdf"
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)
            page.insert_text(
                (72, 96),
                r"PDF题1 中文ABC123 a\times b \sin x a\degree b A\setminus B \frac{a_1}{b^2}",
                fontsize=12,
                fontname="handfont",
                fontfile=str(FONT_PATH),
            )
            doc.save(pdf)
            doc.close()

            with mock.patch("source_extract.extract_pdf_to_markdown", side_effect=MinerUConfigError("MINERU_BASE_URL missing")):
                result = extract_textfileprocess_content(pdf)

        self.assertEqual(result["source"], "pymupdf_pdf_fallback")
        compact_text = re.sub(r"\s+", "", result["text"])
        self.assertIn("PDF题1", compact_text)
        self.assertIn("中文ABC123", compact_text)
        for raw in (r"\degree", r"\setminus"):
            self.assertNotIn(raw, result["text"])
        self.assertIn("$", result["text"])
        for token in (r"a\times b", r"\sin x", "°", "∖", r"\frac{a_1}{b^2}"):
            self.assertIn(token, result["text"])
        debug_text = markdown_render_debug_text(result["text"], FONT_PATH)
        assert_semantic_tokens_preserved(
            self,
            debug_text,
            text_tokens=("PDF题1", "中文ABC123"),
            debug_tokens=("a×b", "sinx", "°", "A∖B", "(a_1)/(b^2)"),
            forbidden_pattern=r"\\|degree|setminus|times|frac",
        )

    def test_multipage_pdf_fallback_preserves_page_tokens_through_render_debug(self):
        import fitz

        from app import extract_textfileprocess_content

        pages = [
            {
                "raw": r"P1-题甲 中文甲A1B2C3 +-*/=<> a\times b \sin x \frac{a_1}{b^2}",
                "text_tokens": ("P1-题甲", "中文甲A1B2C3", "+-*/=<>"),
                "debug_tokens": ("a×b", "sinx", "(a_1)/(b^2)"),
            },
            {
                "raw": r"P2-题乙 中文乙D4E5F6 A\setminus B a\degree b a\diagup b a\leqsim b",
                "text_tokens": ("P2-题乙", "中文乙D4E5F6"),
                "debug_tokens": ("A∖B", "°", "⟋", "⪅"),
            },
            {
                "raw": r"P3-题丙 中文丙G7H8I9 \injlim_{i=1}^{n}x_i x\hphantom{hidden}y \llap{x+y}+\smash{w}",
                "text_tokens": ("P3-题丙", "中文丙G7H8I9"),
                "debug_tokens": ("inj lim", "i=1", "n", "x_i", "xy", "x+y", "w"),
            },
        ]
        with tempfile.TemporaryDirectory() as tmp:
            pdf = Path(tmp) / "multipage_formula.pdf"
            doc = fitz.open()
            for page_data in pages:
                page = doc.new_page(width=595, height=842)
                page.insert_text(
                    (72, 96),
                    page_data["raw"],
                    fontsize=12,
                    fontname="handfont",
                    fontfile=str(FONT_PATH),
                )
            doc.save(pdf)
            doc.close()

            with mock.patch("source_extract.extract_pdf_to_markdown", side_effect=MinerUConfigError("MINERU_BASE_URL missing")):
                result = extract_textfileprocess_content(pdf)

        self.assertEqual(result["source"], "pymupdf_pdf_fallback")
        self.assertIn("$", result["text"])
        for raw in (
            r"\degree",
            r"\setminus",
            r"\diagup",
            r"\leqsim",
            r"\injlim",
            r"\hphantom",
            r"\llap",
            r"\smash",
        ):
            self.assertNotIn(raw, result["text"])
        debug_text = markdown_render_debug_text(result["text"], FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|degree|setminus|diagup|leqsim|injlim|hphantom|llap|smash|hidden")
        for page_data in pages:
            assert_semantic_tokens_preserved(
                self,
                debug_text,
                text_tokens=page_data["text_tokens"],
                debug_tokens=page_data["debug_tokens"],
            )

    def test_source_extract_pdf_uses_text_layer_when_mineru_unavailable(self):
        import fitz

        with tempfile.TemporaryDirectory() as tmp:
            pdf = Path(tmp) / "source_text_layer.pdf"
            doc = fitz.open()
            page = doc.new_page(width=595, height=842)
            page.insert_text(
                (72, 96),
                r"前端PDF 中文ABC123 +-*/=<> a\times b \sin x A\setminus B \frac{a_1}{b^2}",
                fontsize=12,
                fontname="handfont",
                fontfile=str(FONT_PATH),
            )
            doc.save(pdf)
            doc.close()

            with mock.patch("source_extract.extract_pdf_to_markdown", side_effect=MinerUConfigError("MINERU_BASE_URL missing")):
                result = extract_source_to_markdown(pdf)

        self.assertEqual(result["source"], "pymupdf_pdf_fallback")
        self.assertEqual(result["metadata"]["fallback"], "pdf_text_layer")
        self.assertTrue(any("MinerU 尚未配置" in warning for warning in result["warnings"]))
        self.assertTrue(any("PDF 文本层提取" in warning for warning in result["warnings"]))
        markdown = result["markdown"]
        self.assertIn("$", markdown)
        self.assertNotIn(r"\setminus", markdown)
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        assert_semantic_tokens_preserved(
            self,
            debug_text,
            text_tokens=("前端PDF", "中文ABC123", "+-*/=<>"),
            debug_tokens=("a×b", "sinx", "A∖B", "(a_1)/(b^2)"),
            forbidden_pattern=r"\\|setminus|times|frac",
        )

    def test_raw_latex_commands_in_text_are_rendered_as_math(self):
        debug_text = markdown_render_debug_text(
            r"题目 a\equiv b\pmod{n} 结束；因此 \therefore x\ne0，且 y\not\in B。",
            FONT_PATH,
        )
        self.assertNotIn("\\", debug_text)
        for token in ("题目", "a", "≡", "b", "mod", "n", "结束", "∴", "x", "≠", "0", "y", "∉", "B"):
            self.assertIn(token, debug_text)

    def test_bare_additional_common_symbol_commands_are_wrapped_and_rendered(self):
        markdown = r"题目 A\bigcup B 且 a\leqq b，另有 \varpi+\varsigma。"
        normalized = normalize_math_markdown(markdown)
        self.assertNotIn(r"A\bigcup B", normalized.replace(r"$A\bigcup B$", ""))
        self.assertIn(r"$A\bigcup B$", normalized)
        self.assertIn(r"$a\leqq b$", normalized)
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|bigcup|leqq|varpi|varsigma")
        for token in ("题目", "A", "⋃", "B", "且", "≦", "另有", "ϖ", "ς"):
            self.assertIn(token, debug_text)

    def test_bare_ams_symbol_commands_are_wrapped_and_rendered(self):
        markdown = r"题目 a\nleq b 且 A\nsubseteq B，另有 \beth+\Game。"
        normalized = normalize_math_markdown(markdown)
        self.assertIn(r"$a\nleq b$", normalized)
        self.assertIn(r"$A\nsubseteq B$", normalized)
        self.assertIn(r"$\beth+\Game$", normalized)
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|nleq|nsubseteq|beth|Game")
        for token in ("题目", "≰", "且", "⊈", "另有", "ℶ", "⅁"):
            self.assertIn(token, debug_text)

    def test_bare_more_common_symbols_are_wrapped_and_rendered(self):
        markdown = r"题目 a\nparallel b 且 \coprod_{i=1}^{n}A_i，另有 \clubsuit+\natural。"
        normalized = normalize_math_markdown(markdown)
        self.assertIn(r"$a\nparallel b$", normalized)
        self.assertIn(r"$\coprod_{i=1}^{n}A_i$", normalized)
        self.assertIn(r"$\clubsuit+\natural$", normalized)
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|nparallel|coprod|clubsuit|natural")
        for token in ("题目", "∦", "且", "∐", "i=1", "n", "A", "另有", "♣", "♮"):
            self.assertIn(token, debug_text)

    def test_bare_structural_helpers_are_wrapped_and_rendered(self):
        markdown = r"题目 \stackrel{def}{=} 且 \mathrel{R} 结束。"
        normalized = normalize_math_markdown(markdown)
        self.assertIn(r"$\stackrel{def}{=}$", normalized)
        self.assertIn(r"$\mathrel{R}$", normalized)
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|stackrel|mathrel")
        for token in ("题目", "def", "=", "且", "R", "结束"):
            self.assertIn(token, debug_text)

    def test_bare_legacy_buildrel_is_wrapped_and_rendered(self):
        markdown = r"题目 \buildrel def \over = 结束。"
        normalized = normalize_math_markdown(markdown)
        self.assertIn(r"$\overset{def}{=}$", normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 1)
        self.assertFalse(docx_info["has_latex_residuals"])
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|buildrel|over")
        for token in ("题目", "def", "=", "结束"):
            self.assertIn(token, debug_text)

    def test_bare_legacy_dimension_infix_commands_are_wrapped_and_rendered(self):
        markdown = (
            r"题目 a \over b 和 n \choose k，另有 a \above 0pt b "
            r"和 c \overwithdelims() d 以及 n \atopwithdelims[] k 结束。"
        )
        normalized = normalize_math_markdown(markdown)
        for formula in (
            r"$\frac{a}{b}$",
            r"$\binom{n}{k}$",
            r"$\substack{a\\b}$",
            r"$\left(\frac{c}{d}\right)$",
            r"$\left[\substack{n\\k}\right]$",
        ):
            self.assertIn(formula, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 5)
        self.assertFalse(docx_info["has_latex_residuals"])
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|over|choose|above|overwithdelims|atopwithdelims|0pt")
        for token in ("题目", "a", "b", "和", "n", "k", "c", "d", "结束"):
            self.assertIn(token, debug_text)

    def test_bare_operator_delimiter_and_arrow_commands_are_wrapped_as_office_math(self):
        markdown = (
            r"题目 x\iff y，x\sim y，a\circ b+b\star c，"
            r"\lceil x\rceil+\lfloor y\rfloor+\langle v\rangle，"
            r"\sum\limits_{i=1}^{n}x_i+\min_{x\in A}f(x)，"
            r"a\uparrow b+c\downarrow d。"
        )
        normalized = normalize_math_markdown(markdown)
        for formula in (
            r"$x\iff y$",
            r"$x\sim y$",
            r"$a\circ b+b\star c$",
            r"$\lceil x\rceil+\lfloor y\rfloor+\langle v\rangle$",
            r"$\sum_{i=1}^{n}x_i+\min_{x\in A}f(x)$",
            r"$a\uparrow b+c\downarrow d$",
        ):
            self.assertIn(formula, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 6)
        self.assertFalse(docx_info["has_latex_residuals"])
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|iff|sim|circ|star|lceil|rfloor|limits|uparrow|downarrow")
        for token in ("↔", "∼", "∘", "⋆", "⌈", "⌊", "〈", "∑", "min", "↑", "↓"):
            self.assertIn(token, debug_text)

    def test_unsupported_presentation_helpers_rewrite_to_office_math(self):
        markdown = (
            r"题目 \textcolor{red}{x+y}，\cancel{z}+\bcancel{y}+\xcancel{x}+\sout{w}，"
            r"\boxed{a+b}+\fbox{c+d}，\hdotsfor{3}，\raisebox{1ex}{q}，"
            r"\operatornamewithlimits{argmax}_{x} f(x) 结束。"
        )
        normalized = normalize_math_markdown(markdown)
        self.assertNotRegex(
            normalized,
            r"textcolor|cancel|bcancel|xcancel|sout|fbox|hdotsfor|raisebox|operatornamewithlimits",
        )
        for token in ("x+y", "z+y+x+w", r"\boxed{a+b}+\boxed{c+d}", r"\cdots\cdots\cdots", "q", r"\operatorname{argmax}_{x}"):
            self.assertIn(token, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 6)
        self.assertFalse(docx_info["has_latex_residuals"])
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|textcolor|cancel|fbox|hdotsfor|raisebox|operatornamewithlimits")
        for token in ("题目", "x", "y", "z", "w", "a+b", "c+d", "⋯", "q", "argmax", "结束"):
            self.assertIn(token, debug_text)

    def test_font_text_spacing_and_reference_helpers_become_office_math(self):
        markdown = (
            r"题目 \boldsymbol{\alpha}+\boldmath{y}+\cal{F}+\Bbb{R}+\textbf{ABC123}，"
            r"\overline{x}+\underline{y}，"
            r"\textnormal{abc}+\textup{ghi}+\textsl{jkl}+\hbox{hbox}，"
            r"a\quad b+a\thinspace b+\eqref{eq:a}+\ref{r1}+\notag+x。"
        )
        normalized = normalize_math_markdown(markdown)
        self.assertNotRegex(
            normalized,
            r"boldmath|\\cal|\\Bbb|textbf|textnormal|textup|textsl|\\hbox|thinspace|eqref|\\ref|notag",
        )
        for token in (
            r"$\textnormal{abc}+\textup{ghi}+\textsl{jkl}+\hbox{hbox}$",
            r"$a\quad b+a\thinspace b+\eqref{eq:a}+\ref{r1}+\notag+x$",
        ):
            self.assertNotIn(token, normalized)
        for token in (
            r"\boldsymbol{\alpha}",
            "y",
            r"\mathcal{F}",
            r"\mathbb{R}",
            r"\text{ABC123}",
            r"\overline{x}",
            r"\underline{y}",
            r"\text{abc}",
            r"\text{ghi}",
            r"\text{jkl}",
            r"\text{hbox}",
            r"a\quad b+a  b",
            r"(\text{eq:a})",
            r"\text{r1}",
            "+x",
        ):
            self.assertIn(token, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 4)
        self.assertFalse(docx_info["has_latex_residuals"])
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|boldmath|cal|Bbb|textbf|textnormal|\\hbox|thinspace|eqref|notag")
        for token in ("α", "y", "F", "R", "ABC123", "¯x", "y_", "abc", "ghi", "jkl", "hbox", "a", "b", "eq:a", "r1", "x"):
            self.assertIn(token, debug_text)

    def test_bare_matrix_environment_in_text_becomes_display_office_math(self):
        markdown = r"题目 \begin{array}{c|c}\hline a&b\\\cline{1-2}c&d\end{array} 结束。"
        normalized = normalize_math_markdown(markdown)
        self.assertIn("$$", normalized)
        self.assertNotRegex(normalized, r"hline|cline")
        for token in ("题目", r"\begin{array}{c|c}", "a&b", "c&d", "结束"):
            self.assertIn(token, normalized)
        docx_info = inspect_docx_math(editable_docx_bytes(normalized))
        self.assertGreaterEqual(docx_info["office_math_objects"], 1)
        self.assertFalse(docx_info["has_latex_residuals"])
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|begin|array|hline|cline|end")
        for token in ("题目", "a", "b", "c", "d", "结束"):
            self.assertIn(token, debug_text)

    def test_raw_latex_commands_do_not_swallow_adjacent_plain_text(self):
        debug_text = markdown_render_debug_text(r"This is a \frac{1}{2} test.", FONT_PATH)
        self.assertEqual("This is a (1)/(2) test.", debug_text)
        self.assertNotIn("\\frac", debug_text)

    def test_mixed_long_content_debug_text_preserves_all_semantic_tokens(self):
        markdown = (
            r"第1题：中文、English-XYZ789、数字 0.125 和运算符 +-*/=<> 都不能漏。"
            r"裸公式 a\times b、\sin x、a\alpha b、A\cup B。"
            "\n\n"
            r"$$"
            r"\frac{a_1}{b^2}+\sqrt{x}+\sum_{i=1}^{n}x_i+"
            r"\begin{pmatrix}1&2\\3&4\end{pmatrix}"
            r"$$"
            "\n\n"
            r"安全改写 a\degree b，A\setminus B，a\diagup b，a\leqsim b，"
            r"\injlim_{i=1}^{n}x_i，x\hphantom{hidden}y，\llap{x+y}+\smash{w}。"
        )
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        assert_semantic_tokens_preserved(
            self,
            debug_text,
            text_tokens=("第1题", "中文", "English-XYZ789", "0.125", "+-*/=<>"),
            debug_tokens=(
                "a×b",
                "sinx",
                "aαb",
                "A∪B",
                "(a_1)/(b^2)",
                "√",
                "∑",
                "i=1",
                "n",
                "x_i",
                "[[1,2];[3,4]]",
                "°",
                "A∖B",
                "⟋",
                "⪅",
                "inj lim",
                "xy",
                "x+y",
                "w",
            ),
            forbidden_pattern=r"\\|frac|sqrt|sum|begin|pmatrix|degree|setminus|diagup|leqsim|injlim|hphantom|llap|smash|hidden",
        )

    def test_renderer_places_text_on_first_ruled_line_band(self):
        background = Image.new("RGB", (900, 1100), "white")
        top_margin = 70
        line_spacing = 96
        left_margin = 70
        right_margin = 70
        draw = ImageDraw.Draw(background)
        first_rule_y = top_margin + line_spacing
        draw.line((left_margin, first_rule_y, background.width - right_margin, first_rule_y), fill="black")
        font = ImageFont.truetype(str(FONT_PATH), 52)
        config = HandwritingRenderConfig(
            line_spacing=line_spacing,
            font_size=52,
            left_margin=left_margin,
            top_margin=top_margin,
            right_margin=right_margin,
            bottom_margin=70,
            word_spacing=1,
            perturb_x_sigma=0,
            perturb_y_sigma=0,
            ink_depth_sigma=0,
        )
        page = render_markdown_handwriting("横线内文字 $\\frac{1}{2}$", background, font, config)[0]
        ink_bbox = ImageChops.difference(background, page).getbbox()
        self.assertIsNotNone(ink_bbox)
        self.assertGreaterEqual(ink_bbox[1], top_margin)
        self.assertGreaterEqual(ink_bbox[3], first_rule_y - 12)

    def test_jittered_markdown_renderer_keeps_ink_inside_ruled_frame(self):
        background = Image.new("RGB", (760, 520), "white")
        top_margin = 60
        line_spacing = 82
        left_margin = 70
        right_margin = 70
        bottom_margin = 60
        draw = ImageDraw.Draw(background)
        for y in range(top_margin + line_spacing, background.height - bottom_margin + 1, line_spacing):
            draw.line((left_margin, y, background.width - right_margin, y), fill="black")
        font = ImageFont.truetype(str(FONT_PATH), 52)
        config = HandwritingRenderConfig(
            line_spacing=line_spacing,
            font_size=52,
            left_margin=left_margin,
            top_margin=top_margin,
            right_margin=right_margin,
            bottom_margin=bottom_margin,
            word_spacing=1,
            perturb_x_sigma=4,
            perturb_y_sigma=4,
            perturb_theta_sigma=0.04,
            ink_depth_sigma=0,
            seed=17,
        )
        page = render_markdown_handwriting(
            r"边界 $\frac{a_1}{b^2}$ 不漏 123ABCxyz，继续 $x\iff y+\lceil t\rceil$。",
            background,
            font,
            config,
        )[0]
        ink_bbox = ImageChops.difference(background, page).getbbox()
        self.assertIsNotNone(ink_bbox)
        self.assertGreaterEqual(ink_bbox[0], left_margin)
        self.assertGreaterEqual(ink_bbox[1], top_margin)
        self.assertLessEqual(ink_bbox[2], background.width - right_margin)
        self.assertLessEqual(ink_bbox[3], background.height - bottom_margin)

    def test_multipage_jittered_render_keeps_each_page_inside_ruled_frame(self):
        background = Image.new("RGB", (620, 360), "white")
        top_margin = 42
        line_spacing = 58
        left_margin = 48
        right_margin = 48
        bottom_margin = 42
        draw = ImageDraw.Draw(background)
        for y in range(top_margin + line_spacing, background.height - bottom_margin + 1, line_spacing):
            draw.line((left_margin, y, background.width - right_margin, y), fill="black")
        font = ImageFont.truetype(str(FONT_PATH), 42)
        config = HandwritingRenderConfig(
            line_spacing=line_spacing,
            font_size=42,
            left_margin=left_margin,
            top_margin=top_margin,
            right_margin=right_margin,
            bottom_margin=bottom_margin,
            word_spacing=1,
            perturb_x_sigma=3,
            perturb_y_sigma=3,
            perturb_theta_sigma=0.035,
            ink_depth_sigma=0,
            seed=29,
        )
        paragraph = (
            r"第1题 ABC123 中文不漏 $\frac{a_1}{b^2}+\sqrt{x}$，"
            r"集合 A\setminus B，角度 a\degree b，矩阵 "
            r"$\begin{pmatrix}1&2\\3&4\end{pmatrix}$。"
        )
        pages = render_markdown_handwriting("\n\n".join([paragraph] * 10), background, font, config)
        self.assertGreater(len(pages), 1)
        for page in pages:
            ink_bbox = ImageChops.difference(background, page).getbbox()
            self.assertIsNotNone(ink_bbox)
            self.assertGreaterEqual(ink_bbox[0], left_margin)
            self.assertGreaterEqual(ink_bbox[1], top_margin)
            self.assertLessEqual(ink_bbox[2], background.width - right_margin)
            self.assertLessEqual(ink_bbox[3], background.height - bottom_margin)

    def test_plain_text_lines_use_one_ruled_line_even_with_large_font(self):
        background = Image.new("RGB", (520, 360), "white")
        top_margin = 40
        line_spacing = 56
        left_margin = 50
        right_margin = 50
        bottom_margin = 40
        draw = ImageDraw.Draw(background)
        for y in range(top_margin + line_spacing, background.height - bottom_margin + 1, line_spacing):
            draw.line((left_margin, y, background.width - right_margin, y), fill="black")
        font = ImageFont.truetype(str(FONT_PATH), 86)
        config = HandwritingRenderConfig(
            line_spacing=line_spacing,
            font_size=86,
            left_margin=left_margin,
            top_margin=top_margin,
            right_margin=right_margin,
            bottom_margin=bottom_margin,
            word_spacing=1,
            perturb_x_sigma=0,
            perturb_y_sigma=0,
            perturb_theta_sigma=0,
            ink_depth_sigma=0,
        )
        page = render_markdown_handwriting("ABC123\n\nDEF456", background, font, config)[0]
        ink_bbox = ImageChops.difference(background, page).getbbox()
        self.assertIsNotNone(ink_bbox)
        self.assertLessEqual(ink_bbox[3], top_margin + line_spacing * 3)

    def test_markdown_renderer_preserves_line_leading_math_operators(self):
        markdown = "\n".join(
            [
                "- 5 + 3 = 2",
                "+ x - y = z",
                "* a = b",
                r"- $\frac{1}{2}$ + ABC123",
            ]
        )
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("-5+3=2", "+x-y=z", "*a=b", "-(1)/(2)+ABC123"):
            self.assertIn(token, compact_text)

    def test_markdown_renderer_preserves_non_markup_hash_and_star_operators(self):
        markdown = "\n".join(
            [
                "#1 题号 ABC123",
                "#define MAX 10",
                "a ** b = c",
                "### 普通标题仍可去掉 Markdown 标记",
            ]
        )
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        compact_text = re.sub(r"\s+", "", debug_text)
        for token in ("#1题号ABC123", "#defineMAX10", "a**b=c"):
            self.assertIn(token, compact_text)
        self.assertIn("普通标题仍可去掉Markdown标记", compact_text)
        self.assertNotIn("###普通标题", compact_text)

    def test_oversized_inline_formula_wraps_within_available_width(self):
        font = ImageFont.truetype(str(FONT_PATH), 52)
        fonts = FontCache(font)
        formula = "$" + "+".join(f"a_{i}" for i in range(1, 18)) + "$"
        lines = _layout_inline(_text_to_boxes("长公式 " + formula + " 结束", fonts, 52), 360, 1)
        self.assertGreater(len(lines), 1)
        self.assertTrue(all(line.width <= 360 for line in lines))

    def test_oversized_display_formula_stays_within_right_margin(self):
        background = Image.new("RGB", (600, 900), "white")
        font = ImageFont.truetype(str(FONT_PATH), 52)
        config = HandwritingRenderConfig(
            line_spacing=90,
            font_size=52,
            left_margin=70,
            top_margin=70,
            right_margin=70,
            bottom_margin=70,
            word_spacing=1,
            perturb_x_sigma=0,
            perturb_y_sigma=0,
            ink_depth_sigma=0,
        )
        page = render_markdown_handwriting(
            "$$\nabcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ1234567890\n$$",
            background,
            font,
            config,
        )[0]
        ink_bbox = ImageChops.difference(background, page).getbbox()
        self.assertIsNotNone(ink_bbox)
        self.assertLessEqual(ink_bbox[2], background.width - config.right_margin)

    def test_unsplittable_wide_display_formula_scales_within_right_margin(self):
        background = Image.new("RGB", (420, 520), "white")
        font = ImageFont.truetype(str(FONT_PATH), 52)
        config = HandwritingRenderConfig(
            line_spacing=80,
            font_size=52,
            left_margin=50,
            top_margin=50,
            right_margin=50,
            bottom_margin=50,
            word_spacing=1,
            perturb_x_sigma=0,
            perturb_y_sigma=0,
            perturb_theta_sigma=0,
            ink_depth_sigma=0,
        )
        row1 = " & ".join(f"abcdefghij{i}" for i in range(8))
        row2 = " & ".join(f"klmnopqrst{i}" for i in range(8))
        page = render_markdown_handwriting(
            f"$$\n\\begin{{pmatrix}}{row1}\\\\{row2}\\end{{pmatrix}}\n$$",
            background,
            font,
            config,
        )[0]
        ink_bbox = ImageChops.difference(background, page).getbbox()
        self.assertIsNotNone(ink_bbox)
        self.assertGreaterEqual(ink_bbox[0], config.left_margin)
        self.assertLessEqual(ink_bbox[2], background.width - config.right_margin)

    def test_tall_display_formula_stays_above_bottom_margin(self):
        background = Image.new("RGB", (500, 320), "white")
        font = ImageFont.truetype(str(FONT_PATH), 72)
        config = HandwritingRenderConfig(
            line_spacing=80,
            font_size=72,
            left_margin=50,
            top_margin=40,
            right_margin=50,
            bottom_margin=40,
            word_spacing=1,
            perturb_x_sigma=0,
            perturb_y_sigma=0,
            perturb_theta_sigma=0,
            ink_depth_sigma=0,
        )
        page = render_markdown_handwriting(
            r"$$\frac{\frac{a}{b}}{\frac{c}{d}}$$",
            background,
            font,
            config,
        )[0]
        ink_bbox = ImageChops.difference(background, page).getbbox()
        self.assertIsNotNone(ink_bbox)
        self.assertLessEqual(ink_bbox[3], background.height - config.bottom_margin)

    def test_tall_inline_formula_stays_above_bottom_margin(self):
        background = Image.new("RGB", (500, 260), "white")
        font = ImageFont.truetype(str(FONT_PATH), 86)
        config = HandwritingRenderConfig(
            line_spacing=70,
            font_size=86,
            left_margin=50,
            top_margin=40,
            right_margin=50,
            bottom_margin=40,
            word_spacing=1,
            perturb_x_sigma=0,
            perturb_y_sigma=0,
            perturb_theta_sigma=0,
            ink_depth_sigma=0,
        )
        page = render_markdown_handwriting(
            r"内联 $\frac{\frac{a}{b}}{\frac{c}{d}}$ 结束",
            background,
            font,
            config,
        )[0]
        ink_bbox = ImageChops.difference(background, page).getbbox()
        self.assertIsNotNone(ink_bbox)
        self.assertLessEqual(ink_bbox[3], background.height - config.bottom_margin)

    def test_missing_math_symbol_glyphs_use_fallback_font(self):
        font = ImageFont.truetype(str(FONT_PATH), 52)
        self.assertIsNone(font.getmask("↔").getbbox())
        box = TextBox("↔", FontCache(font), 52)
        self.assertIsNotNone(box.font.getmask("↔").getbbox())

    def test_more_common_latex_commands_render_as_math_symbols(self):
        expr = (
            r"\mathcal{F}+\mathfrak{c}+a\equiv b\pmod{n},\quad "
            r"x\perp y,\quad l\parallel m,\quad \angle ABC,\quad "
            r"\therefore x\ne 0,\because y\leq z,\colon,\quad y\not\in B,\quad \|v\|"
        )
        debug_text = latex_to_debug_text(expr, FONT_PATH)
        self.assertNotIn("\\", debug_text)
        self.assertNotRegex(
            debug_text,
            r"mathfrak|equiv|pmod|quad|perp|parallel|angle|therefore|because|colon|not∈",
        )
        for token in ("F", "c", "≡", "mod", "n", "⊥", "∥", "∠", "ABC", "∴", "≠", "∵", "≤", ":", "∉", "B", "‖v‖"):
            self.assertIn(token, debug_text)

    def test_norm_delimiter_commands_render_as_double_bars(self):
        debug_text = latex_to_debug_text(r"\lVert v\rVert+\Vert x\Vert", FONT_PATH)
        self.assertNotRegex(debug_text, r"lVert|rVert|\\")
        self.assertEqual("‖v‖+‖x‖", debug_text)

    def test_limits_commands_attach_scripts_to_big_operator(self):
        font = ImageFont.truetype(str(FONT_PATH), 52)
        box = latex_to_box(r"\sum\limits_{i=1}^{n} x_i", FontCache(font), 52)
        first = box.children[0]
        self.assertIsInstance(first, ScriptBox)
        self.assertEqual(first.base.debug_text(), "∑")
        self.assertEqual(first.sub.debug_text(), "i=1")
        self.assertEqual(first.sup.debug_text(), "n")

    def test_contour_integral_commands_render_as_big_operator_symbols(self):
        markdown = r"题目 \oint_C f(z)\,dz+\oiint_S g\,dS+\oiiint_V h\,dV。"
        normalized = normalize_math_markdown(markdown)
        self.assertIn(r"$\oint_C f(z) dz+\oiint_S g dS+\oiiint_V h dV$", normalized)
        debug_text = markdown_render_debug_text(markdown, FONT_PATH)
        self.assertNotRegex(debug_text, r"\\|oint|oiint|oiiint|operatorname")
        for token in ("题目", "∮", "C", "f", "z", "dz", "∯", "S", "g", "dS", "∰", "V", "h", "dV"):
            self.assertIn(token, debug_text)

    def test_mineru_sanitize_preserves_image_placeholders(self):
        sanitized = sanitize_mineru_markdown("题面\n\n![scan](images/p1.png)\n\n答案")
        self.assertIn("题面", sanitized)
        self.assertIn("[图片:![scan](images/p1.png)]", sanitized)
        self.assertIn("答案", sanitized)

    def test_repair_extracted_markdown_restores_split_pingwen_terms(self):
        markdown = (
            "平 分布由细致平衡 稳 $\\pi_1 p_{12}=\\pi_2 p_{21}$ 得\n\n"
            "平 方程给 稳 $\\pi_n=\\pi_{n-1}p_{n-1}$ ，故若上式有限，\n\n"
            "若级数发散，则不存在平 概率分布。稳\n\n"
            "若 $\\mu_0=\\pi$ 为平 分布，则 稳 $\\mu_n=\\pi$ 。"
            "所以 \\~π 为平稳分布；不可约且有平稳分布，故稳 稳 Y 正常返。\n\n"
            "反之若 X 正常返，则有平 概率； 稳 双随机给出不变测度。\n\n"
            "故平稳分布只稳 能与常数测度成比例；无限状态无法归一化。\n\n"
            "故平 分布只稳 能与常数测度成比例；无限状态无法归一化。\n\n"
            "若 X 可逆，则有可逆平稳分布，由 稳 (1) 知 S 有限。\n\n"
            "若 X 可逆，则有可逆平 分布，由 稳 (1) 知 S 有限。\n\n"
            "有限双随机不可约链的平 分稳布为均匀分布。"
        )
        repaired = repair_extracted_markdown_text(markdown)
        for token in (
            "平稳分布由细致平衡",
            "平稳方程给",
            "平稳概率分布",
            "为平稳分布，则",
            "故 Y 正常返",
            "有平稳概率；双随机",
            "平稳分布只能与常数测度成比例",
            "平稳分布，由 (1)",
            "平稳分布为均匀分布",
        ):
            self.assertIn(token, repaired)
        self.assertNotRegex(
            repaired,
            r"平\s+(?:分布|方程|概率分布)|(?:分布|方程|概率分布)[，。]?稳|故稳|只稳|分稳布|由 稳",
        )

    def test_repair_extracted_markdown_removes_random_process_math_hallucinations(self):
        markdown = (
            "一次 i-循环的平均长度为 $m _ { i i }$ ，其中访问 j 的平均次数为 $e _ { i j \\epsilon }$ 。故\n\n"
            "验：取 j = i 时 $e_{i i} {=} 1$ ，得到恒等式 "
            "$\\pi _ { i } \\overline { { - } } \\pi _ { i _ { \\circ } }$\n\n"
            "设 $N _ { n } { = } \\sum _ { k = 1 } ^ { n } "
            "1 _ { \\{ X _ { k } = j \\} } { \\mathfrak { c } }$ 。从 j 出发，\n\n"
            "设正常返态 i 所在互通类为 $C _ { \\circ }$ 。常返类必闭，"
        )
        repaired = repair_extracted_markdown_text(markdown)

        self.assertIn("$e_{i j}$", repaired)
        self.assertIn("$\\pi_{i} = \\pi_{i}$", repaired)
        self.assertIn("1_{ \\{ X_{k} = j \\} }$ 。从 j 出发", repaired)
        self.assertIn("互通类为 $C$ 。", repaired)
        self.assertNotRegex(repaired, r"\\epsilon|\\overline|\\circ|\\mathfrak")

    def test_random_process_pdf_extraction_restores_key_body_text_when_available(self):
        sample_pdf = Path(__file__).resolve().parents[3] / "随机过程三次作业答案.pdf"
        if not sample_pdf.exists():
            self.skipTest("随机过程三次作业答案.pdf sample is not available")
        with mock.patch("source_extract.extract_pdf_to_markdown", side_effect=MinerUConfigError("MINERU_BASE_URL missing")):
            result = extract_source_to_markdown(sample_pdf)
        markdown = result["markdown"]
        for token in ("平稳分布", "平稳方程", "故若上式有限", "平稳概率分布", "归纳得证"):
            self.assertIn(token, markdown)
        self.assertNotRegex(
            markdown,
            r"平\s+(?:分布|方程|概率分布)|(?:分布|方程|概率分布)[，。]?稳|故稳|只稳|分稳布|由 稳",
        )

    def assert_random_process_debug_matches_source_counts(self, source_text, debug_text, *, require_precise_tokens=False):
        compact_debug = re.sub(r"\s+", "", debug_text)
        debug_for_counts = re.sub(r"第\d+页", "", debug_text)
        source_cjk = re.findall(r"[\u4e00-\u9fff]", source_text)
        debug_cjk = re.findall(r"[\u4e00-\u9fff]", debug_for_counts)

        self.assertFalse(Counter(source_cjk) - Counter(debug_cjk))
        self.assertEqual(Counter(re.findall(r"\d", source_text)), Counter(re.findall(r"\d", debug_for_counts)))
        self.assertNotRegex(
            debug_text,
            r"operatorname|backslash|equin|(?<!Lef)tright|Lefttright|e_ijε|π_i¯|只稳|故稳|由 稳",
        )
        if not require_precise_tokens:
            return
        for token in (
            "平稳分布由细致平衡",
            "状态i表示A瓶中黑球数",
            "若π,ρ是两个不同平稳分布",
            "0常返⇔",
            "0正常返⇔",
            "平稳方程给",
            "若级数发散，则不存在平稳概率分布",
            "平均次数为e_ij",
            "π_i=π_i",
            "设N_n=∑_k=1^n1_{X_k=j}。从j出发",
            "设目标为a=2026",
            "平稳分布只能与常数测度成比例",
            "P为有限阶对称矩阵",
        ):
            self.assertIn(token, compact_debug)

    def test_random_process_pdf_fallback_render_debug_preserves_cjk_and_digit_counts_when_available(self):
        sample_pdf = Path(__file__).resolve().parents[3] / "随机过程三次作业答案.pdf"
        if not sample_pdf.exists():
            self.skipTest("随机过程三次作业答案.pdf sample is not available")

        import fitz

        with fitz.open(sample_pdf) as doc:
            source_text = "\n".join(page.get_text(sort=True) for page in doc)
        with mock.patch("source_extract.extract_pdf_to_markdown", side_effect=MinerUConfigError("MINERU_BASE_URL missing")):
            result = extract_source_to_markdown(sample_pdf)
        debug_text = markdown_render_debug_text(result["markdown"], FONT_PATH)

        self.assert_random_process_debug_matches_source_counts(source_text, debug_text)

    def test_random_process_pdf_live_mineru_render_debug_preserves_counts_when_configured(self):
        sample_pdf = Path(__file__).resolve().parents[3] / "随机过程三次作业答案.pdf"
        if not sample_pdf.exists():
            self.skipTest("随机过程三次作业答案.pdf sample is not available")

        import fitz

        with fitz.open(sample_pdf) as doc:
            source_text = "\n".join(page.get_text(sort=True) for page in doc)
        result = extract_source_to_markdown(sample_pdf)
        if result.get("source") != "mineru":
            self.skipTest("MinerU is not configured; fallback coverage already exercises the sample")

        debug_text = markdown_render_debug_text(result["markdown"], FONT_PATH)
        self.assert_random_process_debug_matches_source_counts(
            source_text, debug_text, require_precise_tokens=True
        )

    def test_random_process_pdf_render_audit_draws_all_debug_content_when_available(self):
        sample_pdf = Path(__file__).resolve().parents[3] / "随机过程三次作业答案.pdf"
        if not sample_pdf.exists():
            self.skipTest("随机过程三次作业答案.pdf sample is not available")

        with mock.patch("source_extract.extract_pdf_to_markdown", side_effect=MinerUConfigError("MINERU_BASE_URL missing")):
            result = extract_source_to_markdown(sample_pdf)
        expected_debug = markdown_render_debug_text(result["markdown"], FONT_PATH)

        background = Image.new("RGB", (900, 1300), "white")
        font = ImageFont.truetype(str(FONT_PATH), 42)
        config = HandwritingRenderConfig(
            line_spacing=86,
            font_size=42,
            left_margin=56,
            top_margin=42,
            right_margin=56,
            bottom_margin=42,
            word_spacing=1,
            perturb_x_sigma=1,
            perturb_y_sigma=1,
            perturb_theta_sigma=0.2,
            ink_depth_sigma=4,
        )
        drawn_debug: list[str] = []
        drawn_bounds: list[tuple[int, int, int, int, int, str]] = []
        pages = render_markdown_handwriting(
            result["markdown"],
            background,
            font,
            config,
            draw_debug_sink=drawn_debug,
            draw_bounds_sink=drawn_bounds,
        )

        self.assertGreater(len(pages), 1)
        self.assertTrue(drawn_debug)
        self.assertTrue(drawn_bounds)
        self.assertEqual(re.sub(r"\s+", "", expected_debug), re.sub(r"\s+", "", "".join(drawn_debug)))
        for page_number, left, top, right, bottom, text in drawn_bounds:
            with self.subTest(text=text[:32]):
                self.assertGreaterEqual(page_number, 1)
                self.assertLessEqual(page_number, len(pages))
                self.assertGreaterEqual(left, 0)
                self.assertGreaterEqual(top, 0)
                self.assertLessEqual(right, background.width)
                self.assertLessEqual(bottom, background.height)
                rendered_crop = pages[page_number - 1].crop((left, top, right, bottom))
                blank_crop = background.crop((left, top, right, bottom))
                self.assertIsNotNone(ImageChops.difference(blank_crop, rendered_crop).getbbox())
        self.assertIsNotNone(ImageChops.difference(background, pages[0]).getbbox())
        self.assertIsNotNone(ImageChops.difference(background, pages[-1]).getbbox())

    def test_markdown_renderer_outputs_nonblank_image_and_docx(self):
        background = Image.new("RGB", (900, 1100), "white")
        font = ImageFont.truetype(str(FONT_PATH), 52)
        config = HandwritingRenderConfig(
            line_spacing=96,
            font_size=52,
            left_margin=70,
            top_margin=70,
            right_margin=70,
            bottom_margin=70,
            word_spacing=1,
            perturb_x_sigma=2,
            perturb_y_sigma=2,
            ink_depth_sigma=8,
        )
        pages = render_markdown_handwriting(
            "测试 $\\frac{1}{2}$\n\n$$\n\\begin{pmatrix}1&2\\\\3&4\\end{pmatrix}\n$$\n",
            background,
            font,
            config,
        )
        self.assertGreaterEqual(len(pages), 1)
        self.assertIsNotNone(ImageChops.difference(background, pages[0]).getbbox())

        with tempfile.TemporaryDirectory() as tmp:
            image_path = Path(tmp) / "page.png"
            pages[0].save(image_path)
            docx_bytes = write_images_to_docx_bytes([image_path])
        self.assertGreater(len(docx_bytes), 1000)

    def test_source_extract_markdown_txt_and_docx(self):
        with tempfile.TemporaryDirectory() as tmp:
            tmp_dir = Path(tmp)
            md = tmp_dir / "source.md"
            md.write_text("# 标题\n\n$x+1$\n", encoding="utf-8")
            self.assertIn("$x+1$", extract_source_to_markdown(md)["markdown"])

            txt = tmp_dir / "source.txt"
            txt.write_text("普通文本", encoding="utf-8")
            self.assertEqual(extract_source_to_markdown(txt)["markdown"].strip(), "普通文本")

            docx = tmp_dir / "source.docx"
            document = Document()
            document.add_paragraph("Word 文本 x + 1")
            document.save(docx)
            extracted = extract_source_to_markdown(docx)["markdown"]
            self.assertIn("Word 文本", extracted)

            formula_docx = tmp_dir / "formula.docx"
            formula_docx.write_bytes(editable_docx_bytes(r"Word 公式 $\frac{a_1}{b^2}+\sum_{i=1}^{n}x_i$ 完成"))
            formula_extracted = extract_source_to_markdown(formula_docx)["markdown"]
            self.assertIn("Word 公式", formula_extracted)
            self.assertIn(r"\frac", formula_extracted)
            self.assertIn(r"\sum", formula_extracted)
            self.assertIn("完成", formula_extracted)

    def test_safe_source_filename_keeps_suffix_for_chinese_names(self):
        self.assertEqual(safe_source_filename("随机过程三次作业答案.pdf", ".pdf"), "source.pdf")
        self.assertEqual(safe_source_filename("作业答案.docx", ".docx", "draft"), "draft.docx")
        self.assertEqual(safe_source_filename("homework.pdf", ".pdf"), "homework.pdf")

    def test_mineru_timeout_error_is_user_facing(self):
        message = user_facing_mineru_error(MinerUExtractionError("MinerU request failed: <urlopen error timed out>"))
        self.assertIn("PDF 识别服务 MinerU 连接超时", message)
        self.assertIn("MINERU_BASE_URL", message)

    def test_pdf_without_mineru_config_has_readable_error(self):
        with tempfile.TemporaryDirectory() as tmp:
            pdf = Path(tmp) / "source.pdf"
            pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
            with mock.patch.dict(os.environ, {}, clear=True):
                with self.assertRaises(MinerUConfigError) as ctx:
                    extract_source_to_markdown(pdf)
            self.assertIn("MINERU_BASE_URL", str(ctx.exception))

    def test_mineru_adapter_with_mocked_service(self):
        zip_buffer = io.BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            zf.writestr("result/full.md", "识别文本 $\\frac{1}{2}$\n")

        with tempfile.TemporaryDirectory() as tmp:
            pdf = Path(tmp) / "source.pdf"
            pdf.write_bytes(b"%PDF-1.4\n%%EOF\n")
            with mock.patch.dict(
                os.environ,
                {
                    "MINERU_BASE_URL": "https://mineru.example",
                    "MINERU_API_TOKEN": "token",
                    "MINERU_PUBLIC_BASE_URL": "https://handwrite.example",
                },
                clear=True,
            ), mock.patch.object(mineru_adapter, "_submit_task", return_value="task-1"), mock.patch.object(
                mineru_adapter, "_poll_task", return_value="https://download.example/result.zip"
            ), mock.patch.object(mineru_adapter, "_download", return_value=zip_buffer.getvalue()):
                result = extract_pdf_to_markdown(pdf)

        self.assertEqual(result["source"], "mineru")
        self.assertIn("识别文本", result["markdown"])


if __name__ == "__main__":
    unittest.main()
