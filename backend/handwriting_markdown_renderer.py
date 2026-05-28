from __future__ import annotations

import math
import random
import re
import tempfile
from contextlib import contextmanager
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable

from docx import Document
from docx.shared import Inches
from markdown_math import normalize_latex_math, normalize_math_markdown
from PIL import Image, ImageDraw, ImageFont


GREEK = {
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "delta": "δ",
    "epsilon": "ε",
    "varepsilon": "ε",
    "zeta": "ζ",
    "eta": "η",
    "theta": "θ",
    "vartheta": "ϑ",
    "iota": "ι",
    "kappa": "κ",
    "lambda": "λ",
    "mu": "μ",
    "nu": "ν",
    "xi": "ξ",
    "pi": "π",
    "rho": "ρ",
    "varrho": "ϱ",
    "sigma": "σ",
    "varsigma": "ς",
    "tau": "τ",
    "upsilon": "υ",
    "varpi": "ϖ",
    "varphi": "φ",
    "phi": "φ",
    "chi": "χ",
    "psi": "ψ",
    "omega": "ω",
    "Gamma": "Γ",
    "Omega": "Ω",
    "Theta": "Θ",
    "Lambda": "Λ",
    "Xi": "Ξ",
    "Pi": "Π",
    "Delta": "Δ",
    "Sigma": "Σ",
    "Phi": "Φ",
    "Psi": "Ψ",
}

SYMBOLS = {
    "le": "≤",
    "leq": "≤",
    "leqslant": "≤",
    "ge": "≥",
    "geq": "≥",
    "geqslant": "≥",
    "neq": "≠",
    "ne": "≠",
    "equiv": "≡",
    "cong": "≅",
    "simeq": "≃",
    "approx": "≈",
    "sim": "∼",
    "infty": "∞",
    "times": "×",
    "cdot": "·",
    "ldots": "…",
    "cdots": "⋯",
    "vdots": "⋮",
    "ddots": "⋱",
    "dots": "…",
    "pm": "±",
    "mp": "∓",
    "to": "→",
    "rightarrow": "→",
    "leftarrow": "←",
    "Rightarrow": "⇒",
    "Leftarrow": "⇐",
    "Longrightarrow": "⟹",
    "Longleftarrow": "⟸",
    "leftrightarrow": "↔",
    "Leftrightarrow": "⇔",
    "Longleftrightarrow": "⟺",
    "iff": "↔",
    "implies": "⇒",
    "mapsto": "↦",
    "hookrightarrow": "↪",
    "hookleftarrow": "↩",
    "twoheadrightarrow": "↠",
    "twoheadleftarrow": "↞",
    "rightsquigarrow": "↝",
    "leftharpoonup": "↼",
    "leftharpoondown": "↽",
    "rightharpoonup": "⇀",
    "rightharpoondown": "⇁",
    "leftrightharpoons": "⇋",
    "rightleftharpoons": "⇌",
    "nearrow": "↗",
    "searrow": "↘",
    "nwarrow": "↖",
    "swarrow": "↙",
    "gets": "←",
    "uparrow": "↑",
    "downarrow": "↓",
    "cap": "∩",
    "cup": "∪",
    "setminus": "∖",
    "subset": "⊂",
    "subseteq": "⊆",
    "supset": "⊃",
    "supseteq": "⊇",
    "smallsetminus": "∖",
    "varnothing": "∅",
    "emptyset": "∅",
    "partial": "∂",
    "nabla": "∇",
    "forall": "∀",
    "exists": "∃",
    "in": "∈",
    "notin": "∉",
    "ni": "∋",
    "mid": "|",
    "vert": "|",
    "lvert": "|",
    "rvert": "|",
    "Vert": "‖",
    "lVert": "‖",
    "rVert": "‖",
    "|": "‖",
    "langle": "〈",
    "rangle": "〉",
    "lceil": "⌈",
    "rceil": "⌉",
    "lfloor": "⌊",
    "rfloor": "⌋",
    "wedge": "∧",
    "land": "∧",
    "vee": "∨",
    "lor": "∨",
    "neg": "¬",
    "lnot": "¬",
    "prec": "≺",
    "preceq": "≼",
    "leqq": "≦",
    "geqq": "≧",
    "leqsim": "⪅",
    "geqsim": "⪆",
    "lesssim": "≲",
    "gtrsim": "≳",
    "lessapprox": "⪅",
    "gtrapprox": "⪆",
    "lessgtr": "≶",
    "gtrless": "≷",
    "succ": "≻",
    "succeq": "≽",
    "ll": "≪",
    "gg": "≫",
    "asymp": "≍",
    "doteq": "≐",
    "sqcup": "⊔",
    "sqcap": "⊓",
    "uplus": "⊎",
    "wr": "≀",
    "amalg": "⨿",
    "oplus": "⊕",
    "otimes": "⊗",
    "models": "⊨",
    "vdash": "⊢",
    "dashv": "⊣",
    "vDash": "⊨",
    "Vdash": "⊩",
    "Vvdash": "⊪",
    "smile": "⌣",
    "frown": "⌢",
    "bowtie": "⋈",
    "perp": "⊥",
    "parallel": "∥",
    "angle": "∠",
    "triangle": "△",
    "therefore": "∴",
    "because": "∵",
    "colon": ":",
    "circ": "∘",
    "star": "⋆",
    "ldotp": ".",
    "cdotp": "·",
    "bullet": "•",
    "diamond": "⋄",
    "Box": "□",
    "square": "□",
    "blacksquare": "■",
    "triangleleft": "◁",
    "triangleright": "▷",
    "propto": "∝",
    "degree": "°",
    "prime": "′",
    "Pr": "P",
    "min": "min",
    "max": "max",
    "sup": "sup",
    "inf": "inf",
    "arg": "arg",
    "det": "det",
    "dim": "dim",
    "ker": "ker",
    "gcd": "gcd",
    "ln": "ln",
    "exp": "exp",
    "sin": "sin",
    "cos": "cos",
    "tan": "tan",
    "arcsin": "arcsin",
    "arccos": "arccos",
    "arctan": "arctan",
    "sinh": "sinh",
    "cosh": "cosh",
    "tanh": "tanh",
    "cot": "cot",
    "sec": "sec",
    "csc": "csc",
    "log": "log",
    "Re": "ℜ",
    "Im": "ℑ",
    "ell": "ℓ",
    "hbar": "ℏ",
    "aleph": "ℵ",
    "wp": "℘",
}

BIG_OPERATORS = {
    "sum": "∑",
    "int": "∫",
    "iint": "∬",
    "iiint": "∭",
    "prod": "∏",
    "bigcup": "⋃",
    "bigcap": "⋂",
    "bigsqcup": "⨆",
    "bigvee": "⋁",
    "bigwedge": "⋀",
    "bigoplus": "⨁",
    "bigotimes": "⨂",
    "lim": "lim",
    "limsup": "lim sup",
    "liminf": "lim inf",
    "injlim": "inj lim",
    "projlim": "proj lim",
}
MATRIX_ENVS = {
    "matrix",
    "smallmatrix",
    "pmatrix",
    "bmatrix",
    "Bmatrix",
    "vmatrix",
    "Vmatrix",
    "cases",
    "aligned",
    "alignedat",
    "alignedat*",
    "align",
    "align*",
    "alignat",
    "alignat*",
    "flalign",
    "flalign*",
    "split",
    "gathered",
    "gather",
    "gather*",
    "multline",
    "multline*",
    "eqnarray",
    "eqnarray*",
    "subarray",
    "array",
}
STYLE_COMMANDS = {"displaystyle", "textstyle", "scriptstyle", "scriptscriptstyle", "limits", "nolimits"}
GROUP_WRAPPERS = {
    "mathrm",
    "mathbf",
    "mathbb",
    "mathscr",
    "mathds",
    "mathcal",
    "mathfrak",
    "mathsf",
    "mathtt",
    "mathit",
    "bm",
    "mathrel",
    "mathbin",
    "mathord",
    "mathopen",
    "mathclose",
    "mathpunct",
    "mathinner",
    "smash",
    "rlap",
    "llap",
    "mathclap",
    "operatorname",
    "text",
    "textrm",
    "boldsymbol",
    "overline",
    "underline",
    "hat",
    "widehat",
    "bar",
    "tilde",
    "widetilde",
}
SIZE_DELIMITERS = {"big", "Big", "bigg", "Bigg", "bigl", "bigr", "Bigl", "Bigr", "biggl", "biggr", "Biggl", "Biggr"}
SPACE_COMMANDS = {
    "quad",
    "qquad",
    "enspace",
    "thinspace",
    "medspace",
    "thickspace",
    "negthinspace",
    "negmedspace",
    "negthickspace",
    "hspace",
    "vspace",
}
LATEX_RESIDUAL_RE = re.compile(
    r"\\(?:frac|sqrt|sum|int|begin|end|left|right|mathbf|mathrm|mathbb|operatorname|textstyle|displaystyle|ldots|cdots|dots)"
)
IMAGE_MD_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
HTML_IMAGE_RE = re.compile(r"<img\b[^>]*>", re.I)
COMMAND_FALLBACKS = {
    "overline": "¯",
    "underline": "_",
    "hat": "^",
    "widehat": "^",
    "^": "^",
    "bar": "¯",
    "tilde": "~",
    "widetilde": "~",
    "~": "~",
    "acute": "´",
    "'": "´",
    "grave": "`",
    "`": "`",
    "breve": "˘",
    "check": "ˇ",
    "mathring": "˚",
    '"': "¨",
    "vec": "→",
    "overrightarrow": "→",
    "overleftarrow": "←",
    "overparen": "⌒",
    "underparen": "⌣",
    "underleftarrow": "_←",
    "underrightarrow": "_→",
    "dot": "·",
    "ddot": "··",
}
DELIMITER_COMMANDS = {
    "{": "{",
    "}": "}",
    "lbrace": "{",
    "rbrace": "}",
    "lparen": "(",
    "rparen": ")",
    "lbrack": "[",
    "rbrack": "]",
    "langle": "〈",
    "rangle": "〉",
    "vert": "|",
    "|": "‖",
    "Vert": "‖",
    "lVert": "‖",
    "rVert": "‖",
    ".": "",
}
FALLBACK_FONT_PATHS = [
    Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf"),
    Path("/System/Library/Fonts/Supplemental/STIXGeneral.otf"),
    Path("/System/Library/Fonts/Supplemental/STIXTwoMath.otf"),
    Path("/System/Library/Fonts/Supplemental/Songti.ttc"),
    Path("/System/Library/Fonts/Symbol.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuSansCondensed.ttf"),
    Path("/usr/share/fonts/truetype/dejavu/DejaVuMathTeXGyre.ttf"),
    Path("/usr/share/fonts/opentype/stix/STIXGeneral.otf"),
    Path("/usr/share/fonts/opentype/stix-word/STIX-Regular.otf"),
    Path("/usr/share/fonts/truetype/noto/NotoSansMath-Regular.ttf"),
]


def _font_supports_text(font: ImageFont.FreeTypeFont, text: str) -> bool:
    for ch in text:
        if ch.isspace():
            continue
        try:
            if font.getmask(ch).getbbox() is None:
                return False
        except Exception:
            return False
    return True


@dataclass
class HandwritingRenderConfig:
    line_spacing: int
    font_size: int
    left_margin: int
    top_margin: int
    right_margin: int
    bottom_margin: int
    word_spacing: int = 0
    line_spacing_sigma: float = 0
    font_size_sigma: float = 0
    word_spacing_sigma: float = 0
    perturb_x_sigma: float = 0
    perturb_y_sigma: float = 0
    perturb_theta_sigma: float = 0.04
    ink_depth_sigma: float = 0
    fill: tuple[int, int, int] = (0, 0, 0)
    seed: int = 14790897


class FontCache:
    def __init__(self, base_font):
        self.base_font = base_font
        self._cache: dict[int, ImageFont.FreeTypeFont] = {}
        self._fallback_cache: dict[tuple[Path, int], ImageFont.FreeTypeFont] = {}
        self._fallback_paths = [path for path in FALLBACK_FONT_PATHS if path.exists()]

    def get(self, size: int):
        size = max(8, int(size))
        if size not in self._cache:
            self._cache[size] = self.base_font.font_variant(size=size)
        return self._cache[size]

    def get_for_text(self, text: str, size: int):
        font = self.get(size)
        if _font_supports_text(font, text):
            return font
        for path in self._fallback_paths:
            key = (path, max(8, int(size)))
            if key not in self._fallback_cache:
                self._fallback_cache[key] = ImageFont.truetype(str(path), size=key[1])
            fallback = self._fallback_cache[key]
            if _font_supports_text(fallback, text):
                return fallback
        return font


class DrawContext:
    def __init__(
        self,
        image: Image.Image,
        draw: ImageDraw.ImageDraw,
        fonts: FontCache,
        config: HandwritingRenderConfig,
        rand: random.Random,
    ):
        self.image = image
        self.draw = draw
        self.fonts = fonts
        self.config = config
        self.rand = rand
        self.depth = 0

    def color(self) -> tuple[int, int, int]:
        delta = self.rand.gauss(0, self.config.ink_depth_sigma) if self.config.ink_depth_sigma else 0
        return tuple(max(0, min(255, int(channel + delta))) for channel in self.config.fill)

    def jitter(self) -> tuple[int, int]:
        return (
            round(self.rand.gauss(0, self.config.perturb_x_sigma)) if self.config.perturb_x_sigma else 0,
            round(self.rand.gauss(0, self.config.perturb_y_sigma)) if self.config.perturb_y_sigma else 0,
        )

    @contextmanager
    def nested(self):
        self.depth += 1
        try:
            yield
        finally:
            self.depth -= 1

    def hand_line(self, x1: int, y1: int, x2: int, y2: int, *, width: int = 1, wobble: float | None = None) -> None:
        wobble = wobble if wobble is not None else max(1.0, self.config.perturb_y_sigma or 1)
        distance = math.hypot(x2 - x1, y2 - y1)
        steps = max(3, min(18, int(distance / 28)))
        points = []
        for i in range(steps + 1):
            t = i / steps
            px = x1 + (x2 - x1) * t
            py = y1 + (y2 - y1) * t
            if i not in {0, steps}:
                px += self.rand.gauss(0, wobble * 0.6)
                py += self.rand.gauss(0, wobble)
            points.append((round(px), round(py)))
        for _ in range(2):
            shifted = [
                (
                    x + round(self.rand.gauss(0, max(0.4, wobble * 0.25))),
                    y + round(self.rand.gauss(0, max(0.4, wobble * 0.25))),
                )
                for x, y in points
            ]
            self.draw.line(shifted, fill=self.color(), width=max(1, width), joint="curve")

    def hand_polyline(self, points: list[tuple[int, int]], *, width: int = 1, wobble: float | None = None) -> None:
        for start, end in zip(points, points[1:]):
            self.hand_line(start[0], start[1], end[0], end[1], width=width, wobble=wobble)


class Box:
    width: int
    height: int
    baseline: int

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        raise NotImplementedError

    def debug_text(self) -> str:
        return ""


class TextBox(Box):
    def __init__(self, text: str, fonts: FontCache, size: int):
        self.text = text
        self.size = size
        self.font = fonts.get_for_text(text or " ", size)
        left, top, right, bottom = self.font.getbbox(text or " ")
        self.width = max(1, right - left)
        ascent, descent = self.font.getmetrics()
        self.baseline = ascent
        self.height = max(1, ascent + descent, bottom - top)

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        dx, dy = ctx.jitter()
        actual_size = self.size
        if ctx.config.font_size_sigma:
            actual_size = max(8, round(ctx.rand.gauss(self.size, ctx.config.font_size_sigma)))
        font = ctx.fonts.get_for_text(self.text or " ", actual_size)
        angle_sigma = max(0.0, float(ctx.config.perturb_theta_sigma or 0.0))
        if self.text.strip() and angle_sigma and ctx.depth > 0:
            bbox = font.getbbox(self.text)
            width = max(1, bbox[2] - bbox[0] + 8)
            height = max(1, bbox[3] - bbox[1] + 8)
            layer = Image.new("RGBA", (width, height), (0, 0, 0, 0))
            layer_draw = ImageDraw.Draw(layer)
            layer_draw.text((4 - bbox[0], 4 - bbox[1]), self.text, fill=ctx.color() + (255,), font=font)
            angle = math.degrees(ctx.rand.gauss(0, angle_sigma * 0.45))
            rotated = layer.rotate(angle, expand=True, resample=Image.Resampling.BICUBIC)
            target = (x + dx - (rotated.width - width) // 2, y + dy - (rotated.height - height) // 2)
            ctx.image.paste(rotated.convert(ctx.image.mode), target, rotated)
        else:
            ctx.draw.text((x + dx, y + dy), self.text, fill=ctx.color(), font=font)

    def debug_text(self) -> str:
        return self.text


class HBox(Box):
    def __init__(self, children: list[Box], gap: int = 0):
        self.children = children
        self.gap = gap
        self.width = sum(child.width for child in children) + gap * max(0, len(children) - 1)
        self.baseline = max((child.baseline for child in children), default=0)
        below = max((child.height - child.baseline for child in children), default=0)
        self.height = self.baseline + below

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        cursor = x
        for child in self.children:
            child.draw(ctx, cursor, y + self.baseline - child.baseline)
            cursor += child.width + self.gap

    def debug_text(self) -> str:
        return "".join(child.debug_text() for child in self.children)


class FractionBox(Box):
    def __init__(self, numerator: Box, denominator: Box, pad: int):
        self.numerator = numerator
        self.denominator = denominator
        self.pad = pad
        self.width = max(numerator.width, denominator.width) + pad * 2
        self.baseline = numerator.height + pad + 2
        self.height = numerator.height + denominator.height + pad * 3 + 2

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        nx = x + (self.width - self.numerator.width) // 2 + ctx.rand.choice([-1, 0, 0, 1])
        dx = x + (self.width - self.denominator.width) // 2 + ctx.rand.choice([-1, 0, 0, 1])
        with ctx.nested():
            self.numerator.draw(ctx, nx, y)
        line_y = y + self.numerator.height + self.pad
        ctx.hand_line(
            x + self.pad // 2,
            line_y + ctx.rand.choice([-1, 0, 0, 1]),
            x + self.width - self.pad // 2,
            line_y + ctx.rand.choice([-1, 0, 0, 1]),
            width=max(2, self.pad // 3),
            wobble=max(1.2, self.pad / 5),
        )
        with ctx.nested():
            self.denominator.draw(ctx, dx, line_y + self.pad)

    def debug_text(self) -> str:
        return f"({self.numerator.debug_text()})/({self.denominator.debug_text()})"


class BinomialBox(Box):
    def __init__(self, upper: Box, lower: Box, size: int, fonts: FontCache):
        self.upper = upper
        self.lower = lower
        self.matrix = MatrixBox([[upper], [lower]], "pmatrix", size, fonts)
        self.width = self.matrix.width
        self.height = self.matrix.height
        self.baseline = self.matrix.baseline

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        self.matrix.draw(ctx, x, y)

    def debug_text(self) -> str:
        return f"C({self.upper.debug_text()},{self.lower.debug_text()})"


class SqrtBox(Box):
    def __init__(self, child: Box, size: int, fonts: FontCache):
        self.child = child
        self.root = TextBox("√", fonts, int(size * 1.15))
        self.pad = max(4, size // 16)
        self.width = self.root.width + child.width + self.pad * 2
        self.baseline = max(self.root.baseline, child.baseline + self.pad)
        self.height = max(self.root.height, child.height + self.pad * 2)

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        child_x = x + self.root.width + self.pad
        child_y = y + self.baseline - self.child.baseline + self.pad // 2
        base_y = child_y + self.child.baseline
        tick_y = base_y - max(2, self.pad)
        root_points = [
            (x + max(1, self.pad // 2), tick_y),
            (x + self.root.width // 3, base_y + max(2, self.pad // 2)),
            (x + self.root.width - max(2, self.pad // 2), child_y + max(1, self.pad // 3)),
            (child_x + self.child.width, child_y + max(1, self.pad // 3) + ctx.rand.choice([-1, 0, 1])),
        ]
        ctx.hand_polyline(root_points, width=max(1, self.pad // 3), wobble=max(1.0, self.pad / 4))
        with ctx.nested():
            self.child.draw(ctx, child_x, child_y)

    def debug_text(self) -> str:
        return f"√({self.child.debug_text()})"


class NthRootBox(Box):
    def __init__(self, index: Box, child: Box, size: int, fonts: FontCache):
        self.index = index
        self.sqrt = SqrtBox(child, size, fonts)
        self.offset_x = max(4, index.width // 2)
        self.offset_y = max(3, index.height // 2)
        self.width = self.offset_x + self.sqrt.width
        self.height = max(index.height, self.offset_y + self.sqrt.height)
        self.baseline = self.offset_y + self.sqrt.baseline

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        with ctx.nested():
            self.index.draw(ctx, x + ctx.rand.choice([-1, 0, 1]), y + ctx.rand.choice([-1, 0, 1]))
        with ctx.nested():
            self.sqrt.draw(ctx, x + self.offset_x, y + self.offset_y)

    def debug_text(self) -> str:
        return f"√[{self.index.debug_text()}]{self.sqrt.child.debug_text()}"


class ScriptBox(Box):
    def __init__(self, base: Box, sup: Box | None, sub: Box | None, limits: bool = False):
        self.base = base
        self.sup = sup
        self.sub = sub
        self.limits = limits
        if limits:
            side = max(sup.width if sup else 0, sub.width if sub else 0)
            self.width = max(base.width, side)
            top = (sup.height + 2) if sup else 0
            bottom = (sub.height + 2) if sub else 0
            self.height = top + base.height + bottom
            self.baseline = top + base.baseline
        else:
            side_width = max(sup.width if sup else 0, sub.width if sub else 0)
            self.width = base.width + side_width
            sup_top = sup.height if sup else 0
            sub_bottom = sub.height if sub else 0
            self.baseline = base.baseline + max(0, sup_top - base.baseline // 2)
            self.height = self.baseline + max(base.height - base.baseline, sub_bottom)

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        if self.limits:
            if self.sup:
                with ctx.nested():
                    self.sup.draw(ctx, x + (self.width - self.sup.width) // 2 + ctx.rand.choice([-2, -1, 0, 1]), y)
            base_y = y + (self.sup.height + 2 if self.sup else 0)
            with ctx.nested():
                self.base.draw(ctx, x + (self.width - self.base.width) // 2, base_y)
            if self.sub:
                with ctx.nested():
                    self.sub.draw(ctx, x + (self.width - self.sub.width) // 2 + ctx.rand.choice([-2, -1, 0, 1]), base_y + self.base.height + 2)
            return
        with ctx.nested():
            self.base.draw(ctx, x, y + self.baseline - self.base.baseline)
        if self.sup:
            with ctx.nested():
                self.sup.draw(ctx, x + self.base.width + ctx.rand.choice([-1, 0, 1]), y + ctx.rand.choice([-1, 0, 1]))
        if self.sub:
            with ctx.nested():
                self.sub.draw(ctx, x + self.base.width + ctx.rand.choice([-1, 0, 1]), y + self.baseline + 2 + ctx.rand.choice([-1, 0, 1]))

    def debug_text(self) -> str:
        text = self.base.debug_text()
        if self.sub:
            text += "_" + self.sub.debug_text()
        if self.sup:
            text += "^" + self.sup.debug_text()
        return text


class MatrixBox(Box):
    def __init__(self, rows: list[list[Box]], env: str, size: int, fonts: FontCache):
        self.rows = rows
        self.env = env
        self.pad = max(8, size // 8)
        columns = max((len(row) for row in rows), default=0)
        self.col_widths = [
            max((row[i].width for row in rows if i < len(row)), default=size // 2)
            for i in range(columns)
        ]
        self.row_heights = [max((cell.height for cell in row), default=size) for row in rows]
        body_width = sum(self.col_widths) + self.pad * max(0, columns - 1)
        body_height = sum(self.row_heights) + self.pad * max(0, len(rows) - 1)
        self.left_delim = TextBox(self._left_delim(), fonts, max(size, int(body_height * 0.75))) if self._left_delim() else None
        self.right_delim = TextBox(self._right_delim(), fonts, max(size, int(body_height * 0.75))) if self._right_delim() else None
        self.width = body_width + (self.left_delim.width if self.left_delim else 0) + (self.right_delim.width if self.right_delim else 0) + self.pad * 2
        self.height = max(body_height, self.left_delim.height if self.left_delim else 0, self.right_delim.height if self.right_delim else 0)
        self.baseline = self.height // 2

    def _left_delim(self) -> str:
        return {"pmatrix": "(", "bmatrix": "[", "Bmatrix": "{", "vmatrix": "|", "Vmatrix": "‖", "cases": "{"}.get(self.env, "")

    def _right_delim(self) -> str:
        return {"pmatrix": ")", "bmatrix": "]", "Bmatrix": "}", "vmatrix": "|", "Vmatrix": "‖"}.get(self.env, "")

    def _draw_delim(self, ctx: DrawContext, delim: str, x: int, y: int, height: int, width: int, *, left: bool) -> None:
        stroke = max(1, self.pad // 4)
        if delim in {"(", ")"}:
            inset = max(2, width // 3)
            outer_x = x + (width - inset if left else inset)
            inner_x = x + (inset if left else width - inset)
            mid_y = y + height // 2
            ctx.hand_polyline(
                [
                    (outer_x, y + stroke),
                    (inner_x, y + height // 4),
                    (inner_x, mid_y),
                    (inner_x, y + height * 3 // 4),
                    (outer_x, y + height - stroke),
                ],
                width=stroke,
                wobble=max(1.0, self.pad / 3),
            )
            return
        if delim in {"[", "]"}:
            edge_x = x + (width - stroke if left else stroke)
            inner_x = x + (stroke if left else width - stroke)
            ctx.hand_line(inner_x, y + stroke, edge_x, y + stroke, width=stroke, wobble=max(1.0, self.pad / 4))
            ctx.hand_line(edge_x, y + stroke, edge_x, y + height - stroke, width=stroke, wobble=max(1.0, self.pad / 4))
            ctx.hand_line(edge_x, y + height - stroke, inner_x, y + height - stroke, width=stroke, wobble=max(1.0, self.pad / 4))
            return
        if delim == "|":
            edge_x = x + width // 2
            ctx.hand_line(edge_x, y + stroke, edge_x + ctx.rand.choice([-1, 0, 1]), y + height - stroke, width=stroke, wobble=max(1.0, self.pad / 4))
            return
        if delim == "{":
            mid_y = y + height // 2
            ctx.hand_polyline(
                [
                    (x + width - stroke, y + stroke),
                    (x + width // 3, y + height // 4),
                    (x + width // 2, mid_y),
                    (x + width // 3, y + height * 3 // 4),
                    (x + width - stroke, y + height - stroke),
                ],
                width=stroke,
                wobble=max(1.0, self.pad / 3),
            )
            return
        TextBox(delim, ctx.fonts, height).draw(ctx, x, y)

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        cursor_x = x
        if self.left_delim:
            self._draw_delim(ctx, self._left_delim(), cursor_x, y, self.height, self.left_delim.width, left=True)
            cursor_x += self.left_delim.width + self.pad // 2
        body_y = y + (self.height - (sum(self.row_heights) + self.pad * max(0, len(self.rows) - 1))) // 2
        row_y = body_y
        for row_index, row in enumerate(self.rows):
            cell_x = cursor_x
            for col_index, cell in enumerate(row):
                with ctx.nested():
                    cell.draw(
                        ctx,
                        cell_x + (self.col_widths[col_index] - cell.width) // 2 + ctx.rand.choice([-1, 0, 0, 1]),
                        row_y + (self.row_heights[row_index] - cell.height) // 2 + ctx.rand.choice([-1, 0, 0, 1]),
                    )
                cell_x += self.col_widths[col_index] + self.pad
            row_y += self.row_heights[row_index] + self.pad
        if self.right_delim:
            self._draw_delim(ctx, self._right_delim(), x + self.width - self.right_delim.width, y, self.height, self.right_delim.width, left=False)

    def debug_text(self) -> str:
        rows = ["[" + ",".join(cell.debug_text() for cell in row) + "]" for row in self.rows]
        return "[" + ";".join(rows) + "]"


class DecoratedBox(Box):
    def __init__(self, child: Box, mark: str, size: int):
        self.child = child
        self.mark = mark
        self.size = size
        self.pad = max(3, size // 12)
        extra_top = max(5, size // 6) if mark in {"¯", "^", "~", "→", "←", "⌒", "´", "`", "¨", "˘", "ˇ", "˚", "·", "··"} else 0
        extra_bottom = max(4, size // 8) if mark in {"_", "⌣", "_→", "_←"} else 0
        self.width = child.width + self.pad * 2
        self.height = child.height + extra_top + extra_bottom
        self.baseline = child.baseline + extra_top

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        child_x = x + self.pad
        child_y = y + (self.baseline - self.child.baseline)
        with ctx.nested():
            self.child.draw(ctx, child_x, child_y)
        stroke = max(1, self.size // 28)
        if self.mark == "¯":
            yy = child_y + max(1, self.pad // 2)
            ctx.hand_line(child_x, yy, child_x + self.child.width, yy + ctx.rand.choice([-1, 0, 1]), width=stroke, wobble=max(1.0, self.pad / 3))
        elif self.mark == "_":
            yy = child_y + self.child.height - max(1, self.pad // 2)
            ctx.hand_line(child_x, yy, child_x + self.child.width, yy + ctx.rand.choice([-1, 0, 1]), width=stroke, wobble=max(1.0, self.pad / 3))
        elif self.mark == "^":
            yy = y + max(1, self.pad // 2)
            mid = child_x + self.child.width // 2
            ctx.hand_polyline([(mid - self.pad, yy + self.pad), (mid, yy), (mid + self.pad, yy + self.pad)], width=stroke, wobble=max(1.0, self.pad / 3))
        elif self.mark == "~":
            yy = y + max(1, self.pad)
            points = [
                (child_x, yy),
                (child_x + self.child.width // 3, yy - self.pad // 2),
                (child_x + self.child.width * 2 // 3, yy + self.pad // 2),
                (child_x + self.child.width, yy),
            ]
            ctx.hand_polyline(points, width=stroke, wobble=max(1.0, self.pad / 3))
        elif self.mark in {"´", "`"}:
            yy = y + max(1, self.pad // 2)
            cx = child_x + self.child.width // 2
            direction = 1 if self.mark == "´" else -1
            ctx.hand_line(cx - direction * self.pad, yy + self.pad, cx + direction * self.pad, yy, width=stroke, wobble=max(1.0, self.pad / 3))
        elif self.mark == "¨":
            yy = y + max(1, self.pad // 2)
            cx = child_x + self.child.width // 2
            for offset in (-self.pad // 2, self.pad // 2):
                ctx.draw.ellipse((cx + offset - stroke, yy - stroke, cx + offset + stroke, yy + stroke), fill=ctx.color())
        elif self.mark in {"˘", "ˇ", "˚"}:
            yy = y + max(1, self.pad // 4)
            cx = child_x + self.child.width // 2
            TextBox(self.mark, ctx.fonts, max(8, int(self.size * 0.55))).draw(ctx, cx - self.pad, yy)
        elif self.mark in {"→", "←"}:
            yy = y + max(1, self.pad)
            self._draw_arrow(ctx, child_x, yy, child_x + self.child.width, self.mark, stroke)
        elif self.mark in {"_→", "_←"}:
            yy = child_y + self.child.height + max(1, self.pad // 2)
            self._draw_arrow(ctx, child_x, yy, child_x + self.child.width, self.mark[-1], stroke)
        elif self.mark == "⌒":
            yy = y + max(1, self.pad)
            mid = child_x + self.child.width // 2
            ctx.hand_polyline([(child_x, yy + self.pad), (mid, yy), (child_x + self.child.width, yy + self.pad)], width=stroke, wobble=max(1.0, self.pad / 3))
        elif self.mark == "⌣":
            yy = child_y + self.child.height + max(1, self.pad // 2)
            mid = child_x + self.child.width // 2
            ctx.hand_polyline([(child_x, yy), (mid, yy + self.pad), (child_x + self.child.width, yy)], width=stroke, wobble=max(1.0, self.pad / 3))
        elif self.mark in {"·", "··"}:
            yy = y + max(1, self.pad // 2)
            cx = child_x + self.child.width // 2
            ctx.draw.ellipse((cx - stroke, yy - stroke, cx + stroke, yy + stroke), fill=ctx.color())
            if self.mark == "··":
                ctx.draw.ellipse((cx + self.pad - stroke, yy - stroke, cx + self.pad + stroke, yy + stroke), fill=ctx.color())

    def _draw_arrow(self, ctx: DrawContext, start_x: int, y: int, end_x: int, direction: str, stroke: int) -> None:
        ctx.hand_line(start_x, y, end_x, y + ctx.rand.choice([-1, 0, 1]), width=stroke, wobble=max(1.0, self.pad / 3))
        if direction == "←":
            ctx.hand_polyline([(start_x + self.pad, y - self.pad // 2), (start_x, y), (start_x + self.pad, y + self.pad // 2)], width=stroke, wobble=max(1.0, self.pad / 3))
        else:
            ctx.hand_polyline([(end_x - self.pad, y - self.pad // 2), (end_x, y), (end_x - self.pad, y + self.pad // 2)], width=stroke, wobble=max(1.0, self.pad / 3))

    def debug_text(self) -> str:
        if self.mark == "_":
            return self.child.debug_text() + "_"
        if self.mark in {"⌣", "_→", "_←"}:
            return self.child.debug_text() + self.mark.replace("_", "")
        return self.mark + self.child.debug_text()


class BoxedBox(Box):
    def __init__(self, child: Box, size: int):
        self.child = child
        self.pad = max(4, size // 10)
        self.width = child.width + self.pad * 2
        self.height = child.height + self.pad * 2
        self.baseline = child.baseline + self.pad

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        stroke = max(1, self.pad // 3)
        ctx.hand_line(x, y, x + self.width, y + ctx.rand.choice([-1, 0, 1]), width=stroke, wobble=max(1.0, self.pad / 3))
        ctx.hand_line(x + self.width, y, x + self.width, y + self.height, width=stroke, wobble=max(1.0, self.pad / 3))
        ctx.hand_line(x + self.width, y + self.height, x, y + self.height + ctx.rand.choice([-1, 0, 1]), width=stroke, wobble=max(1.0, self.pad / 3))
        ctx.hand_line(x, y + self.height, x, y, width=stroke, wobble=max(1.0, self.pad / 3))
        with ctx.nested():
            self.child.draw(ctx, x + self.pad, y + self.pad)

    def debug_text(self) -> str:
        return f"[{self.child.debug_text()}]"


class StrikeBox(Box):
    def __init__(self, child: Box, style: str, size: int):
        self.child = child
        self.style = style
        self.pad = max(2, size // 18)
        self.width = child.width + self.pad * 2
        self.height = child.height + self.pad * 2
        self.baseline = child.baseline + self.pad

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        with ctx.nested():
            self.child.draw(ctx, x + self.pad, y + self.pad)
        stroke = max(1, self.pad // 2)
        left = x + self.pad
        right = x + self.pad + self.child.width
        top = y + self.pad
        bottom = y + self.pad + self.child.height
        mid = top + self.child.height // 2
        if self.style in {"slash", "x"}:
            ctx.hand_line(left, bottom, right, top, width=stroke, wobble=max(1.0, self.pad / 3))
        if self.style in {"backslash", "x"}:
            ctx.hand_line(left, top, right, bottom, width=stroke, wobble=max(1.0, self.pad / 3))
        if self.style == "horizontal":
            ctx.hand_line(left, mid, right, mid + ctx.rand.choice([-1, 0, 1]), width=stroke, wobble=max(1.0, self.pad / 3))

    def debug_text(self) -> str:
        return self.child.debug_text()


class ScaledBox(Box):
    def __init__(self, child: Box, scale: float):
        self.child = child
        self.scale = max(0.01, min(1.0, scale))
        self.width = max(1, int(child.width * self.scale))
        self.height = max(1, int(child.height * self.scale))
        self.baseline = max(1, int(child.baseline * self.scale))

    def draw(self, ctx: DrawContext, x: int, y: int) -> None:
        layer = Image.new("RGBA", (max(1, self.child.width), max(1, self.child.height)), (0, 0, 0, 0))
        layer_ctx = DrawContext(layer, ImageDraw.Draw(layer), ctx.fonts, ctx.config, ctx.rand)
        layer_ctx.depth = ctx.depth
        self.child.draw(layer_ctx, 0, 0)
        resized = layer.resize((self.width, self.height), Image.Resampling.LANCZOS)
        ctx.image.paste(resized.convert(ctx.image.mode), (x, y), resized)

    def debug_text(self) -> str:
        return self.child.debug_text()


class LatexParser:
    def __init__(self, text: str, fonts: FontCache, size: int):
        self.text = text.strip()
        self.fonts = fonts
        self.size = size
        self.pos = 0

    def parse(self) -> Box:
        box = self._parse_until("")
        return box if box.width else TextBox(" ", self.fonts, self.size)

    def _parse_until(self, terminator: str) -> Box:
        children: list[Box] = []
        while self.pos < len(self.text):
            if terminator and self.text.startswith(terminator, self.pos):
                self.pos += len(terminator)
                break
            if terminator == "}" and self.text[self.pos] == "}":
                self.pos += 1
                break
            infix_name = None
            for candidate in ("over", "choose", "atop", "brack", "brace"):
                if self._consume_named_command(candidate):
                    infix_name = candidate
                    break
            if infix_name:
                while children and isinstance(children[-1], TextBox) and children[-1].text.isspace():
                    children.pop()
                self._skip_space()
                numerator = HBox(children, gap=max(0, self.size // 18)) if children else TextBox(" ", self.fonts, self.size)
                denominator = self._parse_until(terminator)
                if infix_name == "over":
                    return FractionBox(numerator, denominator, max(5, self.size // 10))
                if infix_name == "choose":
                    return BinomialBox(numerator, denominator, self.size, self.fonts)
                if infix_name == "brack":
                    return MatrixBox([[numerator], [denominator]], "bmatrix", self.size, self.fonts)
                if infix_name == "brace":
                    return MatrixBox([[numerator], [denominator]], "Bmatrix", self.size, self.fonts)
                return MatrixBox([[numerator], [denominator]], "substack", self.size, self.fonts)
            atom = self._parse_atom()
            atom = self._parse_scripts(atom)
            if atom:
                children.append(atom)
        return HBox(children, gap=max(0, self.size // 18))

    def _parse_atom(self) -> Box:
        ch = self.text[self.pos]
        if ch.isspace():
            self.pos += 1
            return TextBox(" ", self.fonts, self.size // 2)
        if ch == "{":
            self.pos += 1
            return self._parse_until("}")
        if ch == "\\":
            return self._parse_command()
        if ch == "&":
            self.pos += 1
            return TextBox(" ", self.fonts, self.size // 2)
        self.pos += 1
        return TextBox(ch, self.fonts, self.size)

    def _read_command_name(self) -> str:
        self.pos += 1
        start = self.pos
        while self.pos < len(self.text) and self.text[self.pos].isalpha():
            self.pos += 1
        if self.pos == start and self.pos < len(self.text):
            self.pos += 1
        return self.text[start:self.pos]

    def _consume_named_command(self, name: str) -> bool:
        marker = f"\\{name}"
        end = self.pos + len(marker)
        if not self.text.startswith(marker, self.pos):
            return False
        if end < len(self.text) and self.text[end].isalpha():
            return False
        self.pos = end
        return True

    def _skip_optional_star(self) -> None:
        self._skip_space()
        if self.pos < len(self.text) and self.text[self.pos] == "*":
            self.pos += 1

    def _read_group_text(self) -> str:
        self._skip_space()
        if self.pos >= len(self.text) or self.text[self.pos] != "{":
            if self.pos < len(self.text):
                ch = self.text[self.pos]
                self.pos += 1
                return ch
            return ""
        depth = 0
        start = self.pos + 1
        self.pos += 1
        while self.pos < len(self.text):
            ch = self.text[self.pos]
            if ch == "\\" and self.pos + 1 < len(self.text) and self.text[self.pos + 1] in "{}":
                self.pos += 2
                continue
            if ch == "{":
                depth += 1
            elif ch == "}":
                if depth == 0:
                    content = self.text[start:self.pos]
                    self.pos += 1
                    return content
                depth -= 1
            self.pos += 1
        return self.text[start:]

    def _read_raw_group(self) -> str:
        self._skip_space()
        if self.pos >= len(self.text) or self.text[self.pos] != "{":
            return ""
        depth = 0
        start = self.pos
        self.pos += 1
        while self.pos < len(self.text):
            ch = self.text[self.pos]
            if ch == "\\" and self.pos + 1 < len(self.text) and self.text[self.pos + 1] in "{}":
                self.pos += 2
                continue
            if ch == "{":
                depth += 1
            elif ch == "}":
                if depth == 0:
                    self.pos += 1
                    return self.text[start:self.pos]
                depth -= 1
            self.pos += 1
        return self.text[start:]

    def _read_raw_command_arguments(self, limit: int = 2) -> str:
        saved = self.pos
        parts: list[str] = []
        for _ in range(limit):
            raw = self._read_raw_group()
            if not raw:
                break
            parts.append(raw)
        if not parts:
            self.pos = saved
        return "".join(parts)

    def _read_optional_text(self) -> str | None:
        self._skip_space()
        if self.pos >= len(self.text) or self.text[self.pos] != "[":
            return None
        depth = 0
        start = self.pos + 1
        self.pos += 1
        while self.pos < len(self.text):
            ch = self.text[self.pos]
            if ch == "[":
                depth += 1
            elif ch == "]":
                if depth == 0:
                    content = self.text[start:self.pos]
                    self.pos += 1
                    return content
                depth -= 1
            self.pos += 1
        return self.text[start:]

    def _parse_optional_bracket(self, scale: float = 1.0) -> Box | None:
        content = self._read_optional_text()
        if content is None:
            return None
        return LatexParser(content, self.fonts, max(8, int(self.size * scale))).parse()

    def _parse_matrix_content(self, content: str, env: str, scale: float = 0.86) -> MatrixBox:
        rows = []
        for raw_row in re.split(r"\\\\", content):
            cells = [
                LatexParser(cell.strip(), self.fonts, max(8, int(self.size * scale))).parse()
                for cell in raw_row.split("&")
            ]
            if cells:
                rows.append(cells)
        return MatrixBox(rows, env, self.size, self.fonts)

    def _parse_group(self, scale: float = 1.0) -> Box:
        self._skip_space()
        if self.pos < len(self.text) and self.text[self.pos] != "{":
            old_size = self.size
            self.size = max(8, int(self.size * scale))
            try:
                atom = self._parse_atom()
                return self._parse_scripts(atom)
            finally:
                self.size = old_size
        content = self._read_group_text()
        return LatexParser(content, self.fonts, max(8, int(self.size * scale))).parse()

    def _read_delimiter(self) -> str:
        self._skip_space()
        if self.pos >= len(self.text):
            return ""
        if self.text[self.pos] == "\\":
            name = self._read_command_name()
            return DELIMITER_COMMANDS.get(name, SYMBOLS.get(name, name))
        delimiter = self.text[self.pos]
        self.pos += 1
        return "" if delimiter == "." else delimiter

    def _delimiter_text_from_group_content(self, content: str) -> str:
        content = content.strip()
        if not content or content == ".":
            return ""
        if content.startswith("\\"):
            return LatexParser(content, self.fonts, self.size).parse().debug_text()
        return content

    def _parse_unknown_command(self, name: str) -> Box:
        args: list[Box] = []
        while len(args) < 3:
            self._skip_space()
            if self.pos >= len(self.text) or self.text[self.pos] != "{":
                break
            args.append(self._parse_group(0.88))
        if not args:
            return TextBox(name, self.fonts, self.size)
        children: list[Box] = [TextBox(name, self.fonts, max(8, int(self.size * 0.86))), TextBox("(", self.fonts, self.size)]
        for index, arg in enumerate(args):
            if index:
                children.append(TextBox(",", self.fonts, self.size))
            children.append(arg)
        children.append(TextBox(")", self.fonts, self.size))
        return HBox(children, gap=max(0, self.size // 24))

    def _parse_not_command(self) -> Box:
        self._skip_space()
        if self.pos >= len(self.text):
            return TextBox("¬", self.fonts, self.size)
        if self.text[self.pos] == "\\":
            next_name = self._read_command_name()
            negated = {
                "in": "∉",
                "ni": "∌",
                "subset": "⊄",
                "subseteq": "⊈",
                "supset": "⊅",
                "supseteq": "⊉",
                "le": "≰",
                "leq": "≰",
                "ge": "≱",
                "geq": "≱",
                "equiv": "≢",
                "=": "≠",
            }
            return TextBox(negated.get(next_name, "¬" + SYMBOLS.get(next_name, next_name)), self.fonts, self.size)
        ch = self.text[self.pos]
        self.pos += 1
        return TextBox({"=": "≠", "<": "≮", ">": "≯"}.get(ch, "¬" + ch), self.fonts, self.size)

    def _parse_command(self) -> Box:
        name = self._read_command_name()
        if name in {"left", "right", "middle"}:
            delimiter = self._read_delimiter()
            return TextBox(delimiter, self.fonts, int(self.size * 1.05)) if delimiter else TextBox("", self.fonts, self.size)
        if name in STYLE_COMMANDS:
            return TextBox("", self.fonts, self.size)
        if name in SIZE_DELIMITERS:
            return TextBox("", self.fonts, self.size)
        if name in SPACE_COMMANDS:
            self._skip_optional()
            if name in {"hspace", "vspace"}:
                self._read_raw_group()
            return TextBox(" ", self.fonts, self.size // 2)
        if name in {"color", "textcolor"}:
            self._read_group_text()
            self._skip_space()
            return self._parse_group() if self.pos < len(self.text) and self.text[self.pos] == "{" else TextBox("", self.fonts, self.size)
        if name in {"phantom", "hphantom", "vphantom"}:
            self._read_group_text()
            return TextBox("", self.fonts, self.size)
        if name in {"mbox", "hbox"}:
            return TextBox(self._read_group_text(), self.fonts, self.size)
        if name in {"boxed", "fbox"}:
            return BoxedBox(self._parse_group(0.9), self.size)
        if name in {"cancel", "bcancel", "xcancel", "sout"}:
            styles = {"cancel": "slash", "bcancel": "backslash", "xcancel": "x", "sout": "horizontal"}
            return StrikeBox(self._parse_group(0.9), styles[name], self.size)
        if name == "hline":
            return TextBox("", self.fonts, self.size)
        if name == "cline":
            self._read_group_text()
            return TextBox("", self.fonts, self.size)
        if name in {"multicolumn", "multirow"}:
            self._read_group_text()
            self._read_group_text()
            return self._parse_group(0.9)
        if name in COMMAND_FALLBACKS:
            return DecoratedBox(self._parse_group(), COMMAND_FALLBACKS[name], self.size)
        if name == "operatorname":
            self._skip_optional_star()
            return TextBox(self._read_group_text(), self.fonts, self.size)
        if name == "mathop":
            self._skip_optional_star()
            return self._parse_group()
        if name == "substack":
            return self._parse_matrix_content(self._read_group_text(), "substack", 0.82)
        if name == "text":
            return TextBox(self._read_group_text(), self.fonts, self.size)
        if name in GROUP_WRAPPERS:
            return self._parse_group()
        if name in {"frac", "dfrac", "tfrac", "cfrac"}:
            numerator = self._parse_group(0.8)
            denominator = self._parse_group(0.8)
            return FractionBox(numerator, denominator, max(5, self.size // 10))
        if name == "genfrac":
            left = self._delimiter_text_from_group_content(self._read_group_text())
            right = self._delimiter_text_from_group_content(self._read_group_text())
            self._read_group_text()
            self._read_group_text()
            numerator = self._parse_group(0.78)
            denominator = self._parse_group(0.78)
            fraction = FractionBox(numerator, denominator, max(5, self.size // 10))
            children: list[Box] = []
            if left:
                children.append(TextBox(left, self.fonts, self.size))
            children.append(fraction)
            if right:
                children.append(TextBox(right, self.fonts, self.size))
            return HBox(children, gap=max(0, self.size // 24))
        if name in {"binom", "dbinom", "tbinom"}:
            upper = self._parse_group(0.78)
            lower = self._parse_group(0.78)
            return BinomialBox(upper, lower, self.size, self.fonts)
        if name == "sqrt":
            index = self._parse_optional_bracket(0.48)
            child = self._parse_group(0.9)
            return NthRootBox(index, child, self.size, self.fonts) if index else SqrtBox(child, self.size, self.fonts)
        if name in {"xrightarrow", "xleftarrow"}:
            below = self._parse_optional_bracket(0.56)
            above = self._parse_group(0.56)
            arrow = TextBox("→" if name == "xrightarrow" else "←", self.fonts, int(self.size * 1.18))
            return ScriptBox(arrow, above, below, limits=True)
        if name == "overset":
            over = self._parse_group(0.62)
            base = self._parse_group(0.92)
            return ScriptBox(base, over, None, limits=True)
        if name == "stackrel":
            over = self._parse_group(0.62)
            base = self._parse_group(0.92)
            return ScriptBox(base, over, None, limits=True)
        if name == "underset":
            under = self._parse_group(0.62)
            base = self._parse_group(0.92)
            return ScriptBox(base, None, under, limits=True)
        if name == "raisebox":
            self._read_group_text()
            self._skip_optional()
            self._skip_optional()
            return self._parse_group()
        if name == "overbrace":
            child = self._parse_group(0.9)
            return ScriptBox(child, TextBox("⏞", self.fonts, max(8, int(self.size * 0.7))), None, limits=True)
        if name == "underbrace":
            child = self._parse_group(0.9)
            return ScriptBox(child, None, TextBox("⏟", self.fonts, max(8, int(self.size * 0.7))), limits=True)
        if name == "hdotsfor":
            count_text = self._read_group_text().strip()
            count = int(count_text) if count_text.isdigit() else 1
            return TextBox("⋯" * max(1, min(count, 12)), self.fonts, self.size)
        if name in {"pmod", "pod", "mod", "bmod"}:
            content = self._parse_group(0.82) if name in {"pmod", "pod"} else TextBox("mod", self.fonts, max(8, int(self.size * 0.86)))
            if name == "pmod":
                return HBox([TextBox("(mod ", self.fonts, self.size), content, TextBox(")", self.fonts, self.size)])
            if name == "pod":
                return HBox([TextBox("(", self.fonts, self.size), content, TextBox(")", self.fonts, self.size)])
            return content
        if name == "tag":
            self._skip_optional_star()
            return HBox([TextBox("(", self.fonts, self.size), self._parse_group(0.82), TextBox(")", self.fonts, self.size)])
        if name in {"label", "notag", "nonumber"}:
            if name == "label":
                self._read_group_text()
            return TextBox("", self.fonts, self.size)
        if name in {"ref", "eqref"}:
            content = self._read_group_text()
            return TextBox(f"({content})" if name == "eqref" else content, self.fonts, self.size)
        if name == "not":
            return self._parse_not_command()
        if name == "begin":
            env = self._read_group_text().strip()
            return self._parse_environment(env)
        if name in BIG_OPERATORS:
            return TextBox(BIG_OPERATORS[name], self.fonts, int(self.size * (1.35 if name != "lim" else 1.0)))
        if name in GREEK:
            return TextBox(GREEK[name], self.fonts, self.size)
        if name in SYMBOLS:
            return TextBox(SYMBOLS[name], self.fonts, self.size)
        if name in {"\\", ",", ";", ":", "!", " "}:
            return TextBox(" ", self.fonts, self.size // 2)
        if name in {"{", "}", "_", "%", "#", "&"}:
            return TextBox(name, self.fonts, self.size)
        return self._parse_unknown_command(name)

    def _parse_environment(self, env: str) -> Box:
        end_marker = f"\\end{{{env}}}"
        end = self.text.find(end_marker, self.pos)
        if end < 0:
            return TextBox(env, self.fonts, self.size)
        content = self.text[self.pos:end]
        self.pos = end + len(end_marker)
        if env == "array" or env in {"alignedat", "alignedat*", "alignat", "alignat*", "subarray"}:
            content = re.sub(r"^\s*\{[^{}]*\}", "", content, count=1)
        if env not in MATRIX_ENVS:
            return LatexParser(content, self.fonts, self.size).parse()
        return self._parse_matrix_content(content, env)

    def _parse_scripts(self, base: Box) -> Box:
        sup = None
        sub = None
        limits_override: bool | None = None
        while True:
            self._skip_space()
            saved = self.pos
            if self.pos < len(self.text) and self.text[self.pos] == "\\":
                name = self._read_command_name()
                if name == "limits":
                    limits_override = True
                    continue
                if name == "nolimits":
                    limits_override = False
                    continue
                self.pos = saved
            if self.pos >= len(self.text) or self.text[self.pos] not in {"^", "_"}:
                break
            marker = self.text[self.pos]
            self.pos += 1
            script = self._parse_group(0.62)
            if marker == "^":
                sup = script
            else:
                sub = script
        limits = limits_override if limits_override is not None else base.debug_text() in {"∑", "∫", "∬", "∭", "∏", "lim", "lim sup", "lim inf", "inj lim", "proj lim"}
        return ScriptBox(base, sup, sub, limits=limits) if sup or sub else base

    def _skip_optional(self) -> None:
        self._skip_space()
        if self.pos >= len(self.text) or self.text[self.pos] != "[":
            return
        depth = 0
        self.pos += 1
        while self.pos < len(self.text):
            if self.text[self.pos] == "[":
                depth += 1
            elif self.text[self.pos] == "]":
                if depth == 0:
                    self.pos += 1
                    return
                depth -= 1
            self.pos += 1

    def _skip_space(self) -> None:
        while self.pos < len(self.text) and self.text[self.pos].isspace():
            self.pos += 1


def latex_to_box(expr: str, fonts: FontCache, size: int) -> Box:
    return LatexParser(normalize_latex_math(expr), fonts, size).parse()


def latex_to_debug_text(expr: str, font_path: Path | None = None, size: int = 48) -> str:
    font_path = font_path or Path(__file__).resolve().parent / "font_assets" / "神韵英子楷书.ttf"
    font = ImageFont.truetype(str(font_path), size=size)
    return latex_to_box(expr, FontCache(font), size).debug_text()


def _strip_markdown_markup(line: str) -> str:
    line = re.sub(r"^\s{0,3}#{1,6}\s*", "", line)
    line = re.sub(r"^\s*[-*+]\s+", "", line)
    line = re.sub(r"^\s*\d+[.)]\s+", lambda m: m.group(0).strip() + " ", line)
    line = line.replace("**", "").replace("__", "").replace("`", "")
    return line


def _visible_image_marker(match: re.Match[str]) -> str:
    return f"[图片:{match.group(0)}]"


def _visible_html_image_marker(match: re.Match[str]) -> str:
    return f"[图片:{match.group(0)}]"


def _blocks(markdown: str) -> list[tuple[str, str]]:
    lines = normalize_math_markdown(markdown).splitlines()
    blocks: list[tuple[str, str]] = []
    paragraph: list[str] = []
    in_math = False
    math_lines: list[str] = []
    for line in lines:
        stripped = line.strip()
        if stripped == "$$":
            if in_math:
                blocks.append(("math", "\n".join(math_lines)))
                math_lines = []
                in_math = False
            else:
                if paragraph:
                    blocks.append(("text", " ".join(paragraph)))
                    paragraph = []
                in_math = True
            continue
        if in_math:
            math_lines.append(line)
            continue
        if not stripped:
            if paragraph:
                blocks.append(("text", " ".join(paragraph)))
                paragraph = []
            continue
        visible_line = IMAGE_MD_RE.sub(_visible_image_marker, line)
        visible_line = HTML_IMAGE_RE.sub(_visible_html_image_marker, visible_line)
        paragraph.append(_strip_markdown_markup(visible_line))
    if math_lines:
        blocks.append(("math", "\n".join(math_lines)))
    if paragraph:
        blocks.append(("text", " ".join(paragraph)))
    return blocks


INLINE_MATH_RE = re.compile(r"\$([^$\n]+)\$")
RAW_LATEX_COMMAND_RE = re.compile(r"\\(?:[A-Za-z]+|[{}_^~'`\"|])")
TEXT_MATH_BOUNDARY_RE = re.compile(r"[\u4e00-\u9fff，。；：！？、]")
LATEX_CONTEXT_CHARS = set("\\{}[]()_^+-=*/<>.,;:|~'\"!&")


def _is_latex_context_char(ch: str) -> bool:
    if TEXT_MATH_BOUNDARY_RE.match(ch):
        return False
    return ch.isalnum() or ch in LATEX_CONTEXT_CHARS


def _raw_latex_span_end(text: str, end: int) -> int:
    i = end
    brace_depth = 0
    while i < len(text):
        ch = text[i]
        if TEXT_MATH_BOUNDARY_RE.match(ch):
            break
        if ch == "{":
            brace_depth += 1
            i += 1
            continue
        if ch == "}":
            if not brace_depth:
                break
            brace_depth -= 1
            i += 1
            continue
        if ch.isspace():
            if brace_depth:
                i += 1
                continue
            next_nonspace = i
            while next_nonspace < len(text) and text[next_nonspace].isspace():
                next_nonspace += 1
            if next_nonspace < len(text) and text[next_nonspace] == "{":
                i = next_nonspace
                continue
            break
        if ch == "\\":
            i += 1
            if i < len(text) and text[i].isalpha():
                while i < len(text) and text[i].isalpha():
                    i += 1
            elif i < len(text):
                i += 1
            continue
        if ch.isalnum() or ch in LATEX_CONTEXT_CHARS:
            i += 1
            continue
        break
    return i


def _raw_latex_span(text: str, pos: int) -> tuple[int, int] | None:
    match = RAW_LATEX_COMMAND_RE.search(text, pos)
    if not match:
        return None
    start = match.start()
    end = _raw_latex_span_end(text, match.end())
    while start > 0 and _is_latex_context_char(text[start - 1]):
        start -= 1
    while start < match.start() and text[start].isspace():
        start += 1
    while end > match.end() and text[end - 1].isspace():
        end -= 1
    return start, end


def _split_top_level_math(expr: str, separators: set[str]) -> list[str]:
    parts: list[str] = []
    start = 0
    brace_depth = 0
    paren_depth = 0
    bracket_depth = 0
    i = 0
    while i < len(expr):
        if expr.startswith("\\left", i) or expr.startswith("\\right", i):
            i += 5 if expr.startswith("\\left", i) else 6
            continue
        ch = expr[i]
        if ch == "\\":
            i += 1
            while i < len(expr) and expr[i].isalpha():
                i += 1
            continue
        if ch == "{":
            brace_depth += 1
        elif ch == "}":
            brace_depth = max(0, brace_depth - 1)
        elif ch == "(":
            paren_depth += 1
        elif ch == ")":
            paren_depth = max(0, paren_depth - 1)
        elif ch == "[":
            bracket_depth += 1
        elif ch == "]":
            bracket_depth = max(0, bracket_depth - 1)
        elif ch in separators and not brace_depth and not paren_depth and not bracket_depth:
            end = i + 1 if ch in {"=", "+", "-"} else i
            part = expr[start:end].strip()
            if part:
                parts.append(part)
            start = i + 1
        i += 1
    tail = expr[start:].strip()
    if tail:
        parts.append(tail)
    return parts if len(parts) > 1 else [expr.strip()]


def _display_math_lines(expr: str, available_width: int, fonts: FontCache, size: int) -> list[str]:
    expr = normalize_latex_math(expr)
    if not expr:
        return []
    if "\\begin{" in expr:
        return [expr]
    raw_lines = [line.strip() for line in re.split(r"\\\\", expr) if line.strip()]
    lines: list[str] = []
    for raw in raw_lines:
        if latex_to_box(raw, fonts, size).width <= available_width:
            lines.append(raw)
            continue
        split = _split_top_level_math(raw, {"="})
        if len(split) == 1:
            split = _split_top_level_math(raw, {"+", "-"})
        if len(split) == 1:
            split = _split_top_level_math(raw, {","})
        lines.extend(split)
    return lines


def _text_to_boxes(text: str, fonts: FontCache, size: int) -> list[Box]:
    boxes: list[Box] = []
    pos = 0
    for match in INLINE_MATH_RE.finditer(text):
        if match.start() > pos:
            boxes.extend(_plain_text_to_boxes(text[pos:match.start()], fonts, size))
        boxes.append(latex_to_box(match.group(1), fonts, size))
        pos = match.end()
    if pos < len(text):
        boxes.extend(_plain_text_to_boxes(text[pos:], fonts, size))
    return boxes


def _plain_text_to_boxes(text: str, fonts: FontCache, size: int) -> list[Box]:
    boxes: list[Box] = []
    pos = 0
    while pos < len(text):
        span = _raw_latex_span(text, pos)
        if not span:
            boxes.extend(TextBox(ch, fonts, size) for ch in text[pos:])
            break
        start, end = span
        if start < pos:
            start = pos
        if start > pos:
            boxes.extend(TextBox(ch, fonts, size) for ch in text[pos:start])
        boxes.append(latex_to_box(text[start:end], fonts, size))
        pos = end
    return boxes


def markdown_render_debug_text(markdown: str, font_path: Path | None = None, size: int = 48) -> str:
    font_path = font_path or Path(__file__).resolve().parent / "font_assets" / "神韵英子楷书.ttf"
    font = ImageFont.truetype(str(font_path), size=size)
    fonts = FontCache(font)
    parts: list[str] = []
    for kind, content in _blocks(markdown):
        if kind == "math":
            parts.append(latex_to_box(content, fonts, size).debug_text())
        else:
            parts.append("".join(box.debug_text() for box in _text_to_boxes(content, fonts, size)))
    return "\n".join(parts)


def _layout_inline(boxes: list[Box], available_width: int, word_spacing: int) -> list[HBox]:
    layout_boxes: list[Box] = []
    for box in boxes:
        layout_boxes.extend(_split_layout_box(box, available_width))
    lines: list[list[Box]] = []
    current: list[Box] = []
    width = 0
    for box in layout_boxes:
        gap = word_spacing if current else 0
        if current and width + gap + box.width > available_width:
            lines.append(current)
            current = [box]
            width = box.width
        else:
            current.append(box)
            width += gap + box.width
    if current:
        lines.append(current)
    return [HBox(line, gap=word_spacing) for line in lines]


def _split_layout_box(box: Box, available_width: int) -> list[Box]:
    if box.width <= available_width:
        return [box]
    if isinstance(box, HBox) and box.children:
        split: list[Box] = []
        for child in box.children:
            split.extend(_split_layout_box(child, available_width))
        return split
    return [ScaledBox(box, max(1, available_width) / max(1, box.width))]


def render_markdown_handwriting(
    markdown: str,
    background: Image.Image,
    font,
    config: HandwritingRenderConfig,
    progress_hook: Callable[[str, str, int], None] | None = None,
) -> list[Image.Image]:
    markdown = normalize_math_markdown(markdown)
    fonts = FontCache(font)
    rand = random.Random(config.seed)
    available_width = background.width - config.left_margin - config.right_margin
    max_y = background.height - config.bottom_margin
    max_line_height = max(1, max_y - config.top_margin)
    min_formula_size = max(16, min(24, int(config.font_size)))
    pages: list[Image.Image] = []
    page = background.copy()
    draw = ImageDraw.Draw(page)
    ctx = DrawContext(page, draw, fonts, config, rand)
    first_line_y = config.top_margin + config.line_spacing
    baseline_y = first_line_y

    def new_page() -> None:
        nonlocal page, draw, ctx, baseline_y
        pages.append(page)
        page = background.copy()
        draw = ImageDraw.Draw(page)
        ctx = DrawContext(page, draw, fonts, config, rand)
        baseline_y = first_line_y

    def draw_line(line: Box, extra_gap: int = 0, center: bool = False) -> None:
        nonlocal baseline_y
        step_count = max(1, math.ceil((line.height + extra_gap) / max(1, config.line_spacing)))
        line_height = step_count * config.line_spacing
        y = max(config.top_margin, baseline_y - line.baseline)
        if y + line.height > max_y and baseline_y > first_line_y:
            new_page()
            y = max(config.top_margin, baseline_y - line.baseline)
        x = config.left_margin + ((available_width - line.width) // 2 if center and line.width < available_width else 0)
        line.draw(ctx, x, y)
        baseline_y += line_height

    for index, (kind, content) in enumerate(_blocks(markdown), start=1):
        if progress_hook:
            progress_hook("rendering", f"正在处理第 {index} 段", min(90, 45 + index))
        if kind == "math":
            for math_line in _display_math_lines(content, available_width, fonts, config.font_size):
                size = config.font_size
                box = latex_to_box(math_line, fonts, size)
                while (box.width > available_width or box.height > max_line_height) and size > min_formula_size:
                    size = max(min_formula_size, int(size * 0.9))
                    box = latex_to_box(math_line, fonts, size)
                if box.width > available_width or box.height > max_line_height:
                    box = ScaledBox(
                        box,
                        min(
                            max(1, available_width) / max(1, box.width),
                            max(1, max_line_height) / max(1, box.height),
                        ),
                    )
                if box.width > available_width:
                    for wrapped in _layout_inline([box], available_width, config.word_spacing):
                        draw_line(wrapped, extra_gap=config.line_spacing // 2, center=True)
                else:
                    draw_line(box, extra_gap=config.line_spacing // 2, center=True)
            continue
        for line in _layout_inline(_text_to_boxes(content, fonts, config.font_size), available_width, config.word_spacing):
            draw_line(line)

    pages.append(page)
    return pages


def images_to_docx(image_paths: Iterable[Path], output_path: Path, dpi: int = 300) -> Path:
    paths = list(image_paths)
    if not paths:
        raise ValueError("No images to export")
    first = Image.open(paths[0])
    width_in = first.width / dpi
    height_in = first.height / dpi
    first.close()
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    for index, path in enumerate(paths):
        if index:
            doc.add_page_break()
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = 0
        paragraph.paragraph_format.space_before = 0
        run = paragraph.add_run()
        run.add_picture(str(path), width=Inches(width_in))
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path


def write_images_to_docx_bytes(image_paths: Iterable[Path]) -> bytes:
    with tempfile.TemporaryDirectory(prefix="handwriting_docx_") as tmp:
        output = Path(tmp) / "handwriting.docx"
        images_to_docx(image_paths, output)
        return output.read_bytes()
